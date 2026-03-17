import fitz  # PyMuPDF
import pandas as pd
from docx import Document
import re, os, subprocess
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from PIL import Image, ImageTk
import io

class AppExamenesPro:
    def __init__(self, root):
        self.root = root
        self.root.title("Generador de Exámenes")
        ancho = 900  # Aumentado para acomodar el visor
        alto = 800
        pos_x = (self.root.winfo_screenwidth() // 2) - (ancho // 2)
        self.root.geometry(f"{ancho}x{alto}+{pos_x}+15")
        
        # Variables para el visor
        self.pdf_documento = None
        self.total_paginas_pdf = 0
        self.paginas_seleccionadas = set()
        self.miniaturas = []
        self.canvas_items = []
        
        # Variable para guardar el nombre del excel cargado
        self.nombre_archivo_excel = "Pro" 

        self.tab_control = ttk.Notebook(root)
        self.tab1 = ttk.Frame(self.tab_control)
        self.tab2 = ttk.Frame(self.tab_control)
        self.tab_control.add(self.tab1, text=' 1. Extraer preguntas de PDF a Excel ')
        self.tab_control.add(self.tab2, text=' 2. Crear Examen en Word ')
        self.tab_control.pack(expand=1, fill="both")

        self.setup_tab1()
        self.setup_tab2()

    def setup_tab1(self):
        self.ruta_pdf = tk.StringVar()
        self.ruta_excel_dest = tk.StringVar()
        self.nombre_excel_sugerido = ""
        
        tk.Label(self.tab1, text="PASO A: Convertir apuntes en base de datos de preguntas", font=("Arial", 12, "bold")).pack(pady=10)
        
        # Frame superior para selección de PDF
        frame_pdf = tk.LabelFrame(self.tab1, text=" 1. Selecciona el PDF original ", padx=10, pady=10)
        frame_pdf.pack(fill="x", padx=20, pady=5)
        
        tk.Entry(frame_pdf, textvariable=self.ruta_pdf, width=65).pack(side=tk.LEFT, padx=5)
        tk.Button(frame_pdf, text="Buscar...", command=self.seleccionar_pdf).pack(side=tk.LEFT)
        
        self.lbl_info_pdf = tk.Label(self.tab1, text="No se ha cargado ningún PDF", font=("Arial", 9, "italic"), fg="gray")
        self.lbl_info_pdf.pack(pady=2)

        # --- NUEVO: VISOR DE PÁGINAS ---
        frame_visor = tk.LabelFrame(self.tab1, text=" 2. Visor de páginas (Haz clic para seleccionar/deseleccionar)", padx=10, pady=10)
        frame_visor.pack(fill="both", expand=True, padx=20, pady=5)
        
        # Controles del visor
        frame_controles = tk.Frame(frame_visor)
        frame_controles.pack(fill="x", pady=5)
        
        tk.Button(frame_controles, text="Seleccionar todas", command=self.seleccionar_todas).pack(side=tk.LEFT, padx=2)
        tk.Button(frame_controles, text="Deseleccionar todas", command=self.deseleccionar_todas).pack(side=tk.LEFT, padx=2)
        tk.Button(frame_controles, text="Invertir selección", command=self.invertir_seleccion).pack(side=tk.LEFT, padx=2)
        
        self.lbl_seleccionadas = tk.Label(frame_controles, text="Páginas seleccionadas: 0", font=("Arial", 10, "bold"), fg="blue")
        self.lbl_seleccionadas.pack(side=tk.RIGHT, padx=10)
        
        # Canvas con scroll para las miniaturas
        canvas_frame = tk.Frame(frame_visor)
        canvas_frame.pack(fill="both", expand=True)
        
        self.canvas_visor = tk.Canvas(canvas_frame, bg='white')
        scrollbar_y = ttk.Scrollbar(canvas_frame, orient="vertical", command=self.canvas_visor.yview)
        scrollbar_x = ttk.Scrollbar(frame_visor, orient="horizontal", command=self.canvas_visor.xview)
        
        self.frame_miniaturas = tk.Frame(self.canvas_visor, bg='white')
        self.frame_miniaturas.bind("<Configure>", lambda e: self.canvas_visor.configure(scrollregion=self.canvas_visor.bbox("all")))
        
        self.canvas_visor.create_window((0, 0), window=self.frame_miniaturas, anchor="nw")
        self.canvas_visor.configure(yscrollcommand=scrollbar_y.set, xscrollcommand=scrollbar_x.set)
        
        self.canvas_visor.pack(side="left", fill="both", expand=True)
        scrollbar_y.pack(side="right", fill="y")
        scrollbar_x.pack(side="bottom", fill="x")
        
        # Frame para entrada manual
        frame_manual = tk.Frame(frame_visor)
        frame_manual.pack(fill="x", pady=10)
        
        tk.Label(frame_manual, text="O introduce páginas manualmente (ej: 7, 9, 15-20):").pack(side=tk.LEFT, padx=5)
        self.ent_pags = tk.Entry(frame_manual, width=30)
        self.ent_pags.pack(side=tk.LEFT, padx=5)
        tk.Button(frame_manual, text="Aplicar selección manual", command=self.aplicar_seleccion_manual).pack(side=tk.LEFT, padx=5)
        
        # Frame para guardar Excel
        frame_ex = tk.LabelFrame(self.tab1, text=" 3. Guardar Almacén (Excel) como... ", padx=10, pady=10)
        frame_ex.pack(fill="x", padx=20, pady=5)
        
        tk.Entry(frame_ex, textvariable=self.ruta_excel_dest, width=65).pack(side=tk.LEFT, padx=5)
        tk.Button(frame_ex, text="Carpeta...", command=self.seleccionar_destino_excel).pack(side=tk.LEFT)
        
        tk.Button(self.tab1, text="EJECUTAR EXTRACCIÓN", command=self.run_extraccion, 
                 bg="#4CAF50", fg="white", font=("Arial", 11, "bold"), height=2).pack(pady=10)

    def seleccionar_pdf(self):
        f = filedialog.askopenfilename(filetypes=[("PDF", "*.pdf")])
        if f:
            self.ruta_pdf.set(f)
            nombre = os.path.splitext(os.path.basename(f))[0]
            self.nombre_excel_sugerido = f"Banco_Preguntas_{nombre}.xlsx"
            self.ruta_excel_dest.set(self.nombre_excel_sugerido)
            
            try:
                if self.pdf_documento:
                    self.pdf_documento.close()
                
                self.pdf_documento = fitz.open(f)
                self.total_paginas_pdf = len(self.pdf_documento)
                self.lbl_info_pdf.config(text=f"El PDF tiene {self.total_paginas_pdf} páginas.", fg="blue")
                
                # Limpiar selecciones anteriores
                self.paginas_seleccionadas.clear()
                
                # Cargar miniaturas
                self.cargar_miniaturas()
                
            except Exception as e:
                self.lbl_info_pdf.config(text=f"Error al leer el PDF: {e}", fg="red")

    def cargar_miniaturas(self):
        # Limpiar frame de miniaturas
        for widget in self.frame_miniaturas.winfo_children():
            widget.destroy()
        
        self.miniaturas = []
        self.canvas_items = []
        
        # Crear grid de miniaturas (3 columnas)
        fila = 0
        columna = 0
        max_columnas = 3
        
        for i in range(self.total_paginas_pdf):
            # Crear frame para cada página
            frame_pagina = tk.Frame(self.frame_miniaturas, bg='white', relief='raised', borderwidth=1)
            frame_pagina.grid(row=fila, column=columna, padx=5, pady=5, sticky='n')
            
            # Obtener miniatura
            try:
                pagina = self.pdf_documento[i]
                zoom = 0.3  # Factor de zoom para miniaturas
                mat = fitz.Matrix(zoom, zoom)
                pix = pagina.get_pixmap(matrix=mat)
                
                # Convertir a imagen PIL
                img_data = pix.tobytes("ppm")
                img = Image.open(io.BytesIO(img_data))
                img_tk = ImageTk.PhotoImage(img)
                self.miniaturas.append(img_tk)  # Guardar referencia
                
                # Label con la imagen
                lbl_img = tk.Label(frame_pagina, image=img_tk, bg='white')
                lbl_img.pack(padx=2, pady=2)
                
                # Label con número de página
                lbl_num = tk.Label(frame_pagina, text=f"Página {i+1}", bg='white', font=("Arial", 8))
                lbl_num.pack()
                
                # Frame para checkbox
                frame_check = tk.Frame(frame_pagina, bg='white')
                frame_check.pack()
                
                var = tk.BooleanVar()
                var.trace_add("write", lambda *args, idx=i: self.actualizar_seleccion_desde_check(idx))
                
                chk = tk.Checkbutton(frame_check, text="Seleccionar", variable=var, bg='white')
                chk.pack(side=tk.LEFT)
                
                # Guardar referencia
                self.canvas_items.append({
                    'frame': frame_pagina,
                    'var': var,
                    'num': i
                })
                
                # Eventos para selección con clic en la imagen
                lbl_img.bind("<Button-1>", lambda e, idx=i: self.toggle_seleccion(idx))
                lbl_num.bind("<Button-1>", lambda e, idx=i: self.toggle_seleccion(idx))
                frame_pagina.bind("<Button-1>", lambda e, idx=i: self.toggle_seleccion(idx))
                
            except Exception as e:
                print(f"Error cargando página {i+1}: {e}")
            
            # Actualizar grid
            columna += 1
            if columna >= max_columnas:
                columna = 0
                fila += 1

    def toggle_seleccion(self, idx):
        """Alternar selección de una página"""
        if 0 <= idx < len(self.canvas_items):
            item = self.canvas_items[idx]
            # Cambiar el estado del checkbox
            nuevo_valor = not item['var'].get()
            item['var'].set(nuevo_valor)
            
            # Actualizar conjunto de seleccionadas
            if nuevo_valor:
                self.paginas_seleccionadas.add(idx + 1)
            else:
                self.paginas_seleccionadas.discard(idx + 1)
            
            # Actualizar entrada manual y contador
            self.actualizar_entrada_manual()
            self.actualizar_contador_seleccionadas()

    def actualizar_seleccion_desde_check(self, idx):
        """Actualizar cuando se usa el checkbox"""
        if 0 <= idx < len(self.canvas_items):
            item = self.canvas_items[idx]
            if item['var'].get():
                self.paginas_seleccionadas.add(idx + 1)
            else:
                self.paginas_seleccionadas.discard(idx + 1)
            
            self.actualizar_entrada_manual()
            self.actualizar_contador_seleccionadas()

    def actualizar_entrada_manual(self):
        """Actualizar el campo de entrada manual con las páginas seleccionadas"""
        if self.paginas_seleccionadas:
            # Convertir a lista ordenada
            paginas = sorted(list(self.paginas_seleccionadas))
            
            # Crear rangos
            rangos = []
            inicio = paginas[0]
            fin = paginas[0]
            
            for i in range(1, len(paginas)):
                if paginas[i] == fin + 1:
                    fin = paginas[i]
                else:
                    if inicio == fin:
                        rangos.append(str(inicio))
                    else:
                        rangos.append(f"{inicio}-{fin}")
                    inicio = paginas[i]
                    fin = paginas[i]
            
            # Añadir último rango
            if inicio == fin:
                rangos.append(str(inicio))
            else:
                rangos.append(f"{inicio}-{fin}")
            
            self.ent_pags.delete(0, tk.END)
            self.ent_pags.insert(0, ", ".join(rangos))
        else:
            self.ent_pags.delete(0, tk.END)

    def actualizar_contador_seleccionadas(self):
        """Actualizar etiqueta con número de páginas seleccionadas"""
        self.lbl_seleccionadas.config(text=f"Páginas seleccionadas: {len(self.paginas_seleccionadas)}")

    def seleccionar_todas(self):
        """Seleccionar todas las páginas"""
        self.paginas_seleccionadas = set(range(1, self.total_paginas_pdf + 1))
        for item in self.canvas_items:
            item['var'].set(True)
        self.actualizar_entrada_manual()
        self.actualizar_contador_seleccionadas()

    def deseleccionar_todas(self):
        """Deseleccionar todas las páginas"""
        self.paginas_seleccionadas.clear()
        for item in self.canvas_items:
            item['var'].set(False)
        self.actualizar_entrada_manual()
        self.actualizar_contador_seleccionadas()

    def invertir_seleccion(self):
        """Invertir selección actual"""
        for i, item in enumerate(self.canvas_items):
            nuevo_valor = not item['var'].get()
            item['var'].set(nuevo_valor)
            if nuevo_valor:
                self.paginas_seleccionadas.add(i + 1)
            else:
                self.paginas_seleccionadas.discard(i + 1)
        self.actualizar_entrada_manual()
        self.actualizar_contador_seleccionadas()

    def aplicar_seleccion_manual(self):
        """Aplicar selección basada en entrada manual"""
        pags_str = self.ent_pags.get()
        if not pags_str.strip():
            return
        
        # Limpiar selección actual
        self.deseleccionar_todas()
        
        try:
            nuevas_paginas = set()
            for p in pags_str.split(','):
                p = p.strip()
                if '-' in p:
                    ini, fin = map(int, p.split('-'))
                    if ini < 1 or fin > self.total_paginas_pdf:
                        raise ValueError(f"Rango {p} fuera de límites (1-{self.total_paginas_pdf})")
                    nuevas_paginas.update(range(ini, fin + 1))
                else:
                    val = int(p)
                    if val < 1 or val > self.total_paginas_pdf:
                        raise ValueError(f"Página {val} no existe")
                    nuevas_paginas.add(val)
            
            # Actualizar selección
            self.paginas_seleccionadas = nuevas_paginas
            for item in self.canvas_items:
                if (item['num'] + 1) in nuevas_paginas:
                    item['var'].set(True)
                else:
                    item['var'].set(False)
            
            self.actualizar_contador_seleccionadas()
            
        except ValueError as ve:
            messagebox.showerror("Error de Rango", f"Página no válida: {ve}")

    def run_extraccion(self):
        try:
            ruta = self.ruta_pdf.get()
            if not ruta:
                messagebox.showwarning("Atención", "Selecciona un PDF primero.")
                return

            if not self.paginas_seleccionadas:
                messagebox.showwarning("Atención", "Selecciona al menos una página.")
                return

            doc_pdf = fitz.open(ruta)
            indices = [i - 1 for i in sorted(self.paginas_seleccionadas)]  # Convertir a índices base 0

            preguntas = []
            pregunta_actual = ""
            
            # Patrón para detectar preguntas
            patron_pregunta = r'^\d+[\.\)](?!\d)'

            for idx in indices:
                if idx >= len(doc_pdf): continue
                for b in doc_pdf[idx].get_text("dict")["blocks"]:
                    if "lines" not in b: continue
                    for l in b["lines"]:
                        txt = "".join([s["text"] for s in l["spans"]]).strip()
                        bold = any(s["flags"] & 2 or "bold" in s["font"].lower() for s in l["spans"])
                        
                        if bold and re.match(patron_pregunta, txt):
                            if pregunta_actual: preguntas.append(pregunta_actual.strip())
                            pregunta_actual = txt
                        elif pregunta_actual and "entregar" not in txt.lower():
                            pregunta_actual += " " + txt
            
            if pregunta_actual: preguntas.append(pregunta_actual.strip())
            
            if not preguntas:
                messagebox.showwarning("Aviso", "No se detectaron preguntas con negrita y numeración simple.")
                return

            pd.DataFrame({"Pregunta": preguntas}).to_excel(self.ruta_excel_dest.get(), index=False)
            messagebox.showinfo("Éxito", f"Excel creado con {len(preguntas)} preguntas.")
            
        except Exception as e: 
            messagebox.showerror("Error", f"Fallo al extraer: {e}")
        finally:
            if doc_pdf:
                doc_pdf.close()

    def seleccionar_destino_excel(self):
        f = filedialog.asksaveasfilename(initialfile=self.nombre_excel_sugerido, defaultextension=".xlsx")
        if f: self.ruta_excel_dest.set(f)

    def setup_tab2(self):
        # ... (el resto del código del tab2 se mantiene igual)
        self.preguntas_db = []
        self.vars_checks = []
        self.lbl_contador = tk.StringVar(value="Seleccionadas: 0 de 6")
        
        tk.Label(self.tab2, text="PASO B: Seleccionar preguntas y crear examen", font=("Arial", 12, "bold")).pack(pady=10)
        
        frame_config = tk.Frame(self.tab2)
        frame_config.pack(pady=5)
        
        btn_load = tk.Button(frame_config, text="Cargar Almacén (Excel)", command=self.load_excel, bg="#f0f0f0", padx=10)
        btn_load.pack(side=tk.LEFT, padx=10)
        
        tk.Label(frame_config, text="Nº de preguntas en plantilla:").pack(side=tk.LEFT)
        self.ent_max_preguntas = tk.Entry(frame_config, width=5)
        self.ent_max_preguntas.insert(0, "6")
        self.ent_max_preguntas.pack(side=tk.LEFT, padx=5)
        self.ent_max_preguntas.bind("<KeyRelease>", self.actualizar_conteo)

        self.display_conteo = tk.Label(self.tab2, textvariable=self.lbl_contador, font=("Arial", 11, "bold"), fg="blue")
        self.display_conteo.pack(pady=5)

        self.canvas_frame = tk.Frame(self.tab2, bd=2, relief="sunken")
        self.canvas_frame.pack(fill="both", expand=True, padx=20)
        
        self.canvas = tk.Canvas(self.canvas_frame)
        sb = ttk.Scrollbar(self.canvas_frame, orient="vertical", command=self.canvas.yview)
        self.scroll_list = tk.Frame(self.canvas)
        self.scroll_list.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.create_window((0, 0), window=self.scroll_list, anchor="nw")
        self.canvas.configure(yscrollcommand=sb.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")

        self.btn_gen_word = tk.Button(self.tab2, text="GENERAR EXAMEN(WORD)", command=self.create_word, bg="#2196F3", fg="white", font=("Arial", 11, "bold"), state="disabled", height=2)
        self.btn_gen_word.pack(pady=20)

    def load_excel(self):
        f = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx")])
        if f:
            self.nombre_archivo_excel = os.path.splitext(os.path.basename(f))[0]
            
            df = pd.read_excel(f)
            self.preguntas_db = df['Pregunta'].tolist()
            self.actualizar_lista()
            self.btn_gen_word.config(state="normal")

    def actualizar_lista(self):
        for w in self.scroll_list.winfo_children(): w.destroy()
        self.vars_checks = []
        for i, p in enumerate(self.preguntas_db):
            v = tk.BooleanVar()
            v.trace_add("write", self.actualizar_conteo)
            self.vars_checks.append(v)
            texto_chk = f"{i+1}. {p[:120]}..." if len(p) > 120 else f"{i+1}. {p}"
            tk.Checkbutton(self.scroll_list, text=texto_chk, variable=v, justify=tk.LEFT, wraplength=600, anchor="w").pack(fill="x", padx=5, pady=2)

    def actualizar_conteo(self, *args):
        try:
            limite = int(self.ent_max_preguntas.get())
        except:
            limite = 0
        seleccionadas = sum(1 for v in self.vars_checks if v.get())
        self.lbl_contador.set(f"Seleccionadas: {seleccionadas} de {limite}")
        self.display_conteo.config(fg="red" if seleccionadas > limite else "blue")

    def maquetar_texto(self, texto):
        texto = re.sub(r'\?', '?\n', texto)
        texto = re.sub(r'\s+([a-zA-Z]\))', r'\n\1', texto)
        texto = re.sub(r'\s+•', r'\n•', texto)
        texto = re.sub(r'\.\s+([A-Z]\.)', r'.\n\1', texto)
        texto = re.sub(r'\.\s+(_+)', r'.\n\1', texto)
        texto = re.sub(r'\.\s+(\d+)\s+', r'.\n\1 ', texto)
        texto = re.sub(r'\s+(\d+[\.\)])', r'\n\1', texto)
        texto = re.sub(r'\n+', '\n', texto)
        return texto.strip()

    def create_word(self):
        try:
            limite = int(self.ent_max_preguntas.get())
        except:
            messagebox.showerror("Error", "Introduce un número válido.")
            return

        sel = [self.preguntas_db[i] for i, v in enumerate(self.vars_checks) if v.get()]
        if not sel: return
        
        plantilla = filedialog.askopenfilename(title="Elige Plantilla examen.docx", filetypes=[("Word", "*.docx")])
        if not plantilla: return
        
        sugerencia_nombre = f"Examen_{self.nombre_archivo_excel}.docx"
        out = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=sugerencia_nombre)
        
        if not out: return

        try:
            doc = Document(plantilla)
            n = 1
            for p in doc.paragraphs:
                marcador = f"Pregunta {n}."
                if marcador in p.text and n <= len(sel) and n <= limite:
                    texto_limpio = self.maquetar_texto(sel[n-1])
                    p.add_run(f" {texto_limpio}")
                    n += 1
            doc.save(out)
            messagebox.showinfo("¡Éxito!", "Examen guardado.")
            os.startfile(os.path.dirname(out))
        except Exception as e: messagebox.showerror("Error", f"Fallo: {e}")

if __name__ == "__main__":
    root = tk.Tk()
    AppExamenesPro(root)
    root.mainloop()