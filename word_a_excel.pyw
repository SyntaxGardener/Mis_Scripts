import os
import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import threading
import subprocess

class NavajaSuizaDocx:
    def __init__(self, root):
        self.root = root
        self.root.title("Extractor de Tablas e Imágenes")
        
        # Configuración de ventana: Centrada horizontalmente y a 5px del borde superior
        w, h = 550, 450
        px = (self.root.winfo_screenwidth() // 2) - (w // 2)
        self.root.geometry(f"{w}x{h}+{px}+5")
        self.root.configure(bg="#f0f0f0")

        self.doc_path = tk.StringVar()
        self.output_dir = tk.StringVar()
        self.create_widgets()

    def create_widgets(self):
        p = {'padx': 20, 'pady': 8}
        header_font = ("Segoe UI", 12, "bold")
        btn_font = ("Segoe UI", 9, "bold")

        tk.Label(self.root, text="Herramientas de Extracción para Word", font=header_font, bg="#f0f0f0", fg="#333").pack(pady=20)

        # Selección de Archivo
        tk.Label(self.root, text="Documento Word (.docx):", bg="#f0f0f0").pack(anchor="w", **p)
        f1 = tk.Frame(self.root, bg="#f0f0f0"); f1.pack(fill="x", **p)
        tk.Entry(f1, textvariable=self.doc_path).pack(side="left", fill="x", expand=True)
        tk.Button(f1, text="Buscar", command=self.sel_file).pack(side="right", padx=(5,0))

        # Selección de Carpeta
        tk.Label(self.root, text="Carpeta de Destino:", bg="#f0f0f0").pack(anchor="w", **p)
        f2 = tk.Frame(self.root, bg="#f0f0f0"); f2.pack(fill="x", **p)
        tk.Entry(f2, textvariable=self.output_dir).pack(side="left", fill="x", expand=True)
        tk.Button(f2, text="Carpeta", command=self.sel_dir).pack(side="right", padx=(5,0))

        # Botones de Acción
        btn_frame = tk.Frame(self.root, bg="#f0f0f0")
        btn_frame.pack(pady=30)

        self.btn_img = tk.Button(btn_frame, text="EXTRAER IMÁGENES", bg="#e67e22", fg="white", 
                                 font=btn_font, width=22, height=2, command=lambda: self.run("img"))
        self.btn_img.grid(row=0, column=0, padx=10)

        self.btn_tab = tk.Button(btn_frame, text="EXTRAER TABLAS A EXCEL", bg="#27ae60", fg="white", 
                                 font=btn_font, width=22, height=2, command=lambda: self.run("tab"))
        self.btn_tab.grid(row=0, column=1, padx=10)

    def sel_file(self):
        self.doc_path.set(filedialog.askopenfilename(filetypes=[("Word", "*.docx")]))

    def sel_dir(self):
        self.output_dir.set(filedialog.askdirectory())

    def extract_images(self, doc_path, out_dir):
        import zipfile
        # Un .docx es en realidad un archivo comprimido. Las fotos están en word/media/
        with zipfile.ZipFile(doc_path) as z:
            imagenes = [f for f in z.namelist() if f.startswith('word/media/')]
            if not imagenes: return 0
            
            img_dir = os.path.join(out_dir, "Imagenes_Extraidas")
            os.makedirs(img_dir, exist_ok=True)
            
            for img_path in imagenes:
                filename = os.path.basename(img_path)
                with open(os.path.join(img_dir, filename), "wb") as f:
                    f.write(z.read(img_path))
            return len(imagenes)

    def extract_tables(self, doc_path, out_dir):
        doc = Document(doc_path)
        if not doc.tables: return 0
        
        excel_path = os.path.join(out_dir, "Tablas_Extraidas.xlsx")
        with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
            for i, table in enumerate(doc.tables):
                data = []
                for row in table.rows:
                    data.append([cell.text.strip() for cell in row.cells])
                
                df = pd.DataFrame(data)
                # Si la tabla tiene encabezados repetidos por la estructura de Word, los limpiamos
                df.to_excel(writer, sheet_name=f"Tabla_{i+1}", index=False, header=False)
        return len(doc.tables)

    def process(self, mode):
        try:
            doc = self.doc_path.get()
            out = self.output_dir.get()
            
            if mode == "img":
                count = self.extract_images(doc, out)
                msg = f"Se han extraído {count} imágenes."
            else:
                count = self.extract_tables(doc, out)
                msg = f"Se han extraído {count} tablas en un archivo Excel."

            if count > 0:
                messagebox.showinfo("Éxito", msg)
                os.startfile(out) # Abre la carpeta destino automáticamente
            else:
                messagebox.showwarning("Aviso", "No se encontró el contenido solicitado en el documento.")
                
        except Exception as e:
            messagebox.showerror("Error", f"Ocurrió un fallo: {str(e)}")
        finally:
            self.btn_img.config(state="normal")
            self.btn_tab.config(state="normal")

    def run(self, mode):
        if not self.doc_path.get() or not self.output_dir.get():
            return messagebox.showwarning("Atención", "Selecciona el archivo y la carpeta de destino.")
        
        self.btn_img.config(state="disabled")
        self.btn_tab.config(state="disabled")
        threading.Thread(target=self.process, args=(mode,), daemon=True).start()

if __name__ == "__main__":
    root = tk.Tk()
    app = NavajaSuizaDocx(root)
    root.mainloop()