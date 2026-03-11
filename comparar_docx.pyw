#!/usr/bin/env python
# -*- coding: utf-8 -*-

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from docx import Document
from docx.shared import RGBColor
import hashlib
import os
from datetime import datetime
import threading
import re

class ComparadorDOCXAvanzado:
    def __init__(self, root):
        self.root = root
        self.root.title("Comparador de Documentos DOCX")
        ancho_ventana = 1000
        alto_ventana = 800
        distancia_superior = 0

        # USAR 'root' en lugar de 'ventana'
        ancho_pantalla = root.winfo_screenwidth()

        # Calcular la posición X para que esté centrada
        posicion_x = (ancho_pantalla // 2) - (ancho_ventana // 2)

        # Aplicar la geometría a 'root'
        root.geometry(f"{ancho_ventana}x{alto_ventana}+{posicion_x}+{distancia_superior}")
        
        # Variables
        self.archivo1 = tk.StringVar()
        self.archivo2 = tk.StringVar()
        self.ignorar_espacios = tk.BooleanVar(value=False)
        self.ignorar_mayusculas = tk.BooleanVar(value=False)
        self.comparar_formato = tk.BooleanVar(value=True)
        self.comparar_colores = tk.BooleanVar(value=True)
        self.comparar_fuentes = tk.BooleanVar(value=True)
        self.mostrar_solo_diferencias = tk.BooleanVar(value=False)
        
        self.setup_ui()
        
    def setup_ui(self):
        # Estilo
        style = ttk.Style()
        style.theme_use('clam')
        
        # Frame principal
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Título
        titulo = ttk.Label(main_frame, text="📄 Comparador Avanzado de Documentos DOCX", 
                          font=('Arial', 16, 'bold'))
        titulo.pack(pady=10)
        
        # Frame para selección de archivos
        frame_archivos = ttk.LabelFrame(main_frame, text="Seleccionar Archivos", padding="10")
        frame_archivos.pack(fill=tk.X, pady=10)
        
        # Archivo 1
        ttk.Label(frame_archivos, text="Documento 1:").grid(row=0, column=0, sticky=tk.W, pady=5)
        ttk.Entry(frame_archivos, textvariable=self.archivo1, width=60).grid(row=0, column=1, padx=5)
        ttk.Button(frame_archivos, text="Examinar...", 
                  command=lambda: self.seleccionar_archivo(1)).grid(row=0, column=2)
        
        # Archivo 2
        ttk.Label(frame_archivos, text="Documento 2:").grid(row=1, column=0, sticky=tk.W, pady=5)
        ttk.Entry(frame_archivos, textvariable=self.archivo2, width=60).grid(row=1, column=1, padx=5)
        ttk.Button(frame_archivos, text="Examinar...", 
                  command=lambda: self.seleccionar_archivo(2)).grid(row=1, column=2)
        
        # Frame para opciones
        frame_opciones = ttk.LabelFrame(main_frame, text="Opciones de Comparación", padding="10")
        frame_opciones.pack(fill=tk.X, pady=10)
        
        # Crear dos columnas para opciones
        opciones_col1 = ttk.Frame(frame_opciones)
        opciones_col1.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        opciones_col2 = ttk.Frame(frame_opciones)
        opciones_col2.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Columna 1 - Opciones de texto
        ttk.Label(opciones_col1, text="📝 Opciones de Texto:", font=('Arial', 10, 'bold')).pack(anchor=tk.W, pady=5)
        ttk.Checkbutton(opciones_col1, text="Ignorar espacios en blanco", 
                       variable=self.ignorar_espacios).pack(anchor=tk.W, padx=20)
        ttk.Checkbutton(opciones_col1, text="Ignorar MAYÚSCULAS/minúsculas", 
                       variable=self.ignorar_mayusculas).pack(anchor=tk.W, padx=20)
        ttk.Checkbutton(opciones_col1, text="Mostrar solo diferencias", 
                       variable=self.mostrar_solo_diferencias).pack(anchor=tk.W, padx=20)
        
        # Columna 2 - Opciones de formato
        ttk.Label(opciones_col2, text="🎨 Opciones de Formato:", font=('Arial', 10, 'bold')).pack(anchor=tk.W, pady=5)
        ttk.Checkbutton(opciones_col2, text="Comparar formato (negrita/cursiva)", 
                       variable=self.comparar_formato).pack(anchor=tk.W, padx=20)
        ttk.Checkbutton(opciones_col2, text="Comparar colores de texto", 
                       variable=self.comparar_colores).pack(anchor=tk.W, padx=20)
        ttk.Checkbutton(opciones_col2, text="Comparar tipos de fuente", 
                       variable=self.comparar_fuentes).pack(anchor=tk.W, padx=20)
        
        # Botones de acción
        frame_botones = ttk.Frame(main_frame)
        frame_botones.pack(pady=10)
        
        ttk.Button(frame_botones, text="🔍 Comparar Documentos", 
                  command=self.iniciar_comparacion, width=20).pack(side=tk.LEFT, padx=5)
        ttk.Button(frame_botones, text="🔄 Intercambiar Archivos", 
                  command=self.intercambiar_archivos, width=20).pack(side=tk.LEFT, padx=5)
        ttk.Button(frame_botones, text="📊 Exportar Resultados", 
                  command=self.exportar_resultados, width=20).pack(side=tk.LEFT, padx=5)
        ttk.Button(frame_botones, text="🧹 Limpiar Resultados", 
                  command=self.limpiar_resultados, width=20).pack(side=tk.LEFT, padx=5)
        
        # Barra de progreso
        self.progreso = ttk.Progressbar(main_frame, mode='indeterminate')
        self.progreso.pack(fill=tk.X, pady=5)
        
        # Área de resultados
        frame_resultados = ttk.LabelFrame(main_frame, text="Resultados de la Comparación", padding="10")
        frame_resultados.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # Crear un widget de texto con scroll
        self.texto_resultados = scrolledtext.ScrolledText(
            frame_resultados, 
            wrap=tk.WORD, 
            width=80, 
            height=20,
            font=('Consolas', 10)
        )
        self.texto_resultados.pack(fill=tk.BOTH, expand=True)
        
        # Configurar etiquetas de color para resultados
        self.texto_resultados.tag_config('exito', foreground='green', font=('Arial', 10, 'bold'))
        self.texto_resultados.tag_config('error', foreground='red', font=('Arial', 10, 'bold'))
        self.texto_resultados.tag_config('info', foreground='blue')
        self.texto_resultados.tag_config('advertencia', foreground='orange', font=('Arial', 10, 'bold'))
        self.texto_resultados.tag_config('diff_texto', background='#fff3cd', foreground='#856404')  # Amarillo para diferencias de texto
        self.texto_resultados.tag_config('diff_formato', background='#d4edda', foreground='#155724')  # Verde para diferencias de formato
        self.texto_resultados.tag_config('diff_color', background='#cce5ff', foreground='#004085')  # Azul para diferencias de color
        self.texto_resultados.tag_config('titulo_seccion', font=('Arial', 11, 'bold'), foreground='#6c757d')
        
        # Barra de estado
        self.status_bar = ttk.Label(main_frame, text="Listo", relief=tk.SUNKEN)
        self.status_bar.pack(fill=tk.X, pady=5)
        
    def seleccionar_archivo(self, num):
        archivo = filedialog.askopenfilename(
            title=f"Seleccionar Documento {num}",
            filetypes=[("Documentos Word", "*.docx"), ("Todos los archivos", "*.*")]
        )
        if archivo:
            if num == 1:
                self.archivo1.set(archivo)
            else:
                self.archivo2.set(archivo)
    
    def intercambiar_archivos(self):
        temp = self.archivo1.get()
        self.archivo1.set(self.archivo2.get())
        self.archivo2.set(temp)
        self.mostrar_mensaje("✅ Archivos intercambiados")
    
    def limpiar_resultados(self):
        self.texto_resultados.delete(1.0, tk.END)
        self.status_bar.config(text="Resultados limpiados")
    
    def iniciar_comparacion(self):
        if not self.archivo1.get() or not self.archivo2.get():
            messagebox.showwarning("Advertencia", "Por favor seleccione ambos archivos")
            return
        
        # Limpiar resultados anteriores
        self.texto_resultados.delete(1.0, tk.END)
        self.progreso.start()
        self.status_bar.config(text="Comparando documentos...")
        
        # Ejecutar comparación en un hilo separado
        thread = threading.Thread(target=self.comparar_documentos)
        thread.daemon = True
        thread.start()
    
    def comparar_documentos(self):
        try:
            archivo1 = self.archivo1.get()
            archivo2 = self.archivo2.get()
            
            # Verificar que existen
            if not os.path.exists(archivo1) or not os.path.exists(archivo2):
                self.root.after(0, self.mostrar_error, "Uno o ambos archivos no existen")
                return
            
            # Mostrar información básica
            self.root.after(0, self.insertar_texto, f"{'='*70}\n", 'titulo_seccion')
            self.root.after(0, self.insertar_texto, f"📄 COMPARACIÓN DE DOCUMENTOS\n", 'info')
            self.root.after(0, self.insertar_texto, f"{'='*70}\n", 'titulo_seccion')
            self.root.after(0, self.insertar_texto, f"\n📁 Documento 1: {os.path.basename(archivo1)}\n", None)
            self.root.after(0, self.insertar_texto, f"📁 Documento 2: {os.path.basename(archivo2)}\n", None)
            self.root.after(0, self.insertar_texto, f"🕐 Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n", None)
            
            # Mostrar opciones activas
            opciones_activas = []
            if self.ignorar_espacios.get(): opciones_activas.append("Ignorar espacios")
            if self.ignorar_mayusculas.get(): opciones_activas.append("Ignorar MAYÚSC/minúsc")
            if self.comparar_formato.get(): opciones_activas.append("Comparar formato")
            if self.comparar_colores.get(): opciones_activas.append("Comparar colores")
            if self.comparar_fuentes.get(): opciones_activas.append("Comparar fuentes")
            
            self.root.after(0, self.insertar_texto, f"⚙️ Opciones: {', '.join(opciones_activas) if opciones_activas else 'Ninguna'}\n", None)
            self.root.after(0, self.insertar_texto, f"\n{'='*70}\n\n", 'titulo_seccion')
            
            # Comparar tamaños
            tam1 = os.path.getsize(archivo1)
            tam2 = os.path.getsize(archivo2)
            self.root.after(0, self.insertar_texto, f"📊 TAMAÑOS:\n", 'info')
            self.root.after(0, self.insertar_texto, f"   DOC1: {self.formato_tamaño(tam1)}\n", None)
            self.root.after(0, self.insertar_texto, f"   DOC2: {self.formato_tamaño(tam2)}\n", None)
            
            if tam1 != tam2:
                self.root.after(0, self.insertar_texto, f"   ⚠️ Diferencia: {self.formato_tamaño(abs(tam1 - tam2))}\n", 'advertencia')
            
            # Comparar hashes
            hash1 = self.calcular_hash(archivo1)
            hash2 = self.calcular_hash(archivo2)
            
            self.root.after(0, self.insertar_texto, f"\n🔒 HASH MD5:\n", 'info')
            self.root.after(0, self.insertar_texto, f"   DOC1: {hash1[:16]}...\n", None)
            self.root.after(0, self.insertar_texto, f"   DOC2: {hash2[:16]}...\n", None)
            
            if hash1 == hash2:
                self.root.after(0, self.insertar_texto, f"\n✅ ¡Los archivos son IDÉNTICOS (mismo hash)!\n", 'exito')
                self.root.after(0, self.status_bar.config, {"text": "Comparación completada - Archivos idénticos"})
                self.root.after(0, self.progreso.stop)
                return
            
            # Comparar contenido DOCX
            doc1 = Document(archivo1)
            doc2 = Document(archivo2)
            
            # Comparar párrafos con formato
            self.root.after(0, self.insertar_texto, f"\n📝 ANÁLISIS DE CONTENIDO:\n", 'info')
            self.comparar_parrafos_con_formato(doc1, doc2)
            
            # Resumen final
            self.root.after(0, self.insertar_texto, f"\n{'='*70}\n", 'titulo_seccion')
            self.root.after(0, self.insertar_texto, f"✅ Comparación completada\n", 'exito')
            self.root.after(0, self.status_bar.config, {"text": "Comparación completada"})
            
        except Exception as e:
            self.root.after(0, self.mostrar_error, f"Error durante la comparación: {str(e)}")
        finally:
            self.root.after(0, self.progreso.stop)
    
    def comparar_parrafos_con_formato(self, doc1, doc2):
        """Compara párrafos incluyendo formato"""
        parrafos1 = list(doc1.paragraphs)
        parrafos2 = list(doc2.paragraphs)
        
        self.root.after(0, self.insertar_texto, f"   Total párrafos DOC1: {len(parrafos1)}\n", None)
        self.root.after(0, self.insertar_texto, f"   Total párrafos DOC2: {len(parrafos2)}\n", None)
        
        if len(parrafos1) != len(parrafos2):
            self.root.after(0, self.insertar_texto, f"   ⚠️ Diferente número de párrafos: {len(parrafos1)} vs {len(parrafos2)}\n", 'advertencia')
        
        # Comparar párrafo por párrafo
        diferencias_texto = 0
        diferencias_formato = 0
        diferencias_color = 0
        
        max_parrafos = max(len(parrafos1), len(parrafos2))
        
        for i in range(max_parrafos):
            if i < len(parrafos1) and i < len(parrafos2):
                p1 = parrafos1[i]
                p2 = parrafos2[i]
                
                # Comparar texto
                texto1 = p1.text
                texto2 = p2.text
                
                if self.ignorar_espacios.get():
                    texto1 = re.sub(r'\s+', ' ', texto1).strip()
                    texto2 = re.sub(r'\s+', ' ', texto2).strip()
                
                if self.ignorar_mayusculas.get():
                    texto1 = texto1.lower()
                    texto2 = texto2.lower()
                
                if texto1 != texto2:
                    diferencias_texto += 1
                    if not self.mostrar_solo_diferencias.get() or diferencias_texto <= 10:
                        self.root.after(0, self.mostrar_diferencia_texto, i+1, p1.text, p2.text)
                
                # Comparar formato si está activado
                if self.comparar_formato.get() and p1.text and p2.text:
                    diff_formato = self.comparar_formato_runs(p1, p2)
                    if diff_formato:
                        diferencias_formato += 1
                        if not self.mostrar_solo_diferencias.get() or diferencias_formato <= 10:
                            self.root.after(0, self.mostrar_diferencia_formato, i+1, diff_formato)
            
            elif i < len(parrafos1):
                # Párrafo solo en documento 1
                self.root.after(0, self.insertar_texto, f"\n📌 Párrafo {i+1} (solo en DOC1):\n", 'advertencia')
                self.root.after(0, self.insertar_texto, f"   {parrafos1[i].text[:100]}\n", None)
            else:
                # Párrafo solo en documento 2
                self.root.after(0, self.insertar_texto, f"\n📌 Párrafo {i+1} (solo en DOC2):\n", 'advertencia')
                self.root.after(0, self.insertar_texto, f"   {parrafos2[i].text[:100]}\n", None)
        
        # Mostrar resumen de diferencias
        self.root.after(0, self.insertar_texto, f"\n📊 RESUMEN DE DIFERENCIAS:\n", 'info')
        self.root.after(0, self.insertar_texto, f"   🔤 Diferencias de texto: {diferencias_texto}\n", None)
        if self.comparar_formato.get():
            self.root.after(0, self.insertar_texto, f"   ✨ Diferencias de formato: {diferencias_formato}\n", None)
        if self.comparar_colores.get():
            self.root.after(0, self.insertar_texto, f"   🎨 Diferencias de color: {diferencias_color}\n", None)
    
    def comparar_formato_runs(self, parrafo1, parrafo2):
        """Compara el formato de los runs en dos párrafos"""
        diferencias = []
        
        runs1 = list(parrafo1.runs)
        runs2 = list(parrafo2.runs)
        
        for j, (run1, run2) in enumerate(zip(runs1, runs2)):
            if run1.text.strip() and run2.text.strip():
                # Comparar negrita
                if run1.bold != run2.bold:
                    diferencias.append(f"      Run {j+1}: Negrita {run1.bold} vs {run2.bold}")
                
                # Comparar cursiva
                if run1.italic != run2.italic:
                    diferencias.append(f"      Run {j+1}: Cursiva {run1.italic} vs {run2.italic}")
                
                # Comparar subrayado
                if run1.underline != run2.underline:
                    diferencias.append(f"      Run {j+1}: Subrayado {run1.underline} vs {run2.underline}")
                
                # Comparar color si está activado
                if self.comparar_colores.get():
                    color1 = run1.font.color.rgb if run1.font.color else None
                    color2 = run2.font.color.rgb if run2.font.color else None
                    if color1 != color2:
                        diferencias.append(f"      Run {j+1}: Color diferente")
                
                # Comparar fuente si está activado
                if self.comparar_fuentes.get():
                    fuente1 = run1.font.name
                    fuente2 = run2.font.name
                    if fuente1 != fuente2:
                        diferencias.append(f"      Run {j+1}: Fuente {fuente1} vs {fuente2}")
        
        return diferencias
    
    def mostrar_diferencia_texto(self, num_parrafo, texto1, texto2):
        """Muestra diferencia de texto formateada"""
        self.root.after(0, self.insertar_texto, f"\n📌 Diferencia en Párrafo {num_parrafo} (TEXTO):\n", 'diff_texto')
        self.root.after(0, self.insertar_texto, f"   DOC1: {texto1[:150]}{'...' if len(texto1) > 150 else ''}\n", None)
        self.root.after(0, self.insertar_texto, f"   DOC2: {texto2[:150]}{'...' if len(texto2) > 150 else ''}\n", None)
    
    def mostrar_diferencia_formato(self, num_parrafo, diferencias):
        """Muestra diferencias de formato"""
        self.root.after(0, self.insertar_texto, f"\n✨ Diferencia en Párrafo {num_parrafo} (FORMATO):\n", 'diff_formato')
        for diff in diferencias[:3]:  # Mostrar solo primeras 3 diferencias
            self.root.after(0, self.insertar_texto, f"{diff}\n", None)
    
    def calcular_hash(self, archivo):
        hash_md5 = hashlib.md5()
        with open(archivo, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def formato_tamaño(self, bytes):
        for unidad in ['B', 'KB', 'MB', 'GB']:
            if bytes < 1024.0:
                return f"{bytes:.1f} {unidad}"
            bytes /= 1024.0
        return f"{bytes:.1f} TB"
    
    def insertar_texto(self, texto, tag=None):
        self.texto_resultados.insert(tk.END, texto, tag)
        self.texto_resultados.see(tk.END)
        self.root.update_idletasks()
    
    def mostrar_error(self, mensaje):
        self.insertar_texto(f"\n❌ ERROR: {mensaje}\n", 'error')
        messagebox.showerror("Error", mensaje)
        self.status_bar.config(text="Error en la comparación")
    
    def mostrar_mensaje(self, mensaje):
        messagebox.showinfo("Información", mensaje)
    
    def exportar_resultados(self):
        if not self.texto_resultados.get(1.0, tk.END).strip():
            messagebox.showwarning("Advertencia", "No hay resultados para exportar")
            return
        
        archivo = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Archivos de texto", "*.txt"), ("Todos los archivos", "*.*")]
        )
        
        if archivo:
            try:
                with open(archivo, 'w', encoding='utf-8') as f:
                    f.write(self.texto_resultados.get(1.0, tk.END))
                messagebox.showinfo("Éxito", f"Resultados exportados a:\n{archivo}")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo exportar: {str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = ComparadorDOCXAvanzado(root)
    root.mainloop()