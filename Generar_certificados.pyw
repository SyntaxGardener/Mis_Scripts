# -*- coding: utf-8 -*-
import os, re, unicodedata
import pandas as pd
from docx import Document
from datetime import datetime
from dateutil import parser
import tkinter as tk
from tkinter import filedialog, messagebox
import threading

# --- CONFIGURACIÓN DE NEGRIFAS ---
MESES = {1:"enero",2:"febrero",3:"marzo",4:"abril",5:"mayo",6:"junio",
         7:"julio",8:"agosto",9:"septiembre",10:"octubre",11:"noviembre",12:"diciembre"}

NEG_CAMPOS = ["nombre", "dni", "media"]
NEG_FRASES = ["Educación Secundaria Obligatoria"]

# --- LÓGICA DE PROCESAMIENTO  ---
def normalizar(t):
    if not isinstance(t,str): t=str(t)
    t=t.strip().lower()
    return ''.join(c for c in unicodedata.normalize('NFD',t) if unicodedata.category(c)!='Mn')

def fecha_castellano(v):
    try:
        if isinstance(v,(int,float)): f = datetime(1899,12,30) + pd.to_timedelta(v,"D")
        else: f = parser.parse(str(v), dayfirst=True)
        return f"{f.day} de {MESES[f.month]} de {f.year}"
    except: return str(v)

def calcular_segmentos(texto_original, sexo, reemplazos):
    """
    Analiza el texto ORIGINAL del párrafo (antes de sustituir nada) y devuelve
    una lista de segmentos a reemplazar: (inicio, fin, texto_nuevo, forzar_negrita).
    Trabajar sobre las posiciones del texto original (en vez de sobre el texto
    ya sustituido) es lo que evita el desajuste de negritas cuando el texto
    insertado tiene una longitud distinta al marcador o a la palabra original.
    """
    s = str(sexo).strip().upper()
    es_m = "M" in s
    cambios = {
        r"D\./Dña\.": "D." if es_m else "Dña.",
        r"nacido/a": "nacido" if es_m else "nacida",
        r"del/de la": "del" if es_m else "de la",
        r"interesado/a": "interesado" if es_m else "interesada",
        r"alumno/a": "alumno" if es_m else "alumna",
        r"el/la": "el" if es_m else "la"
    }

    candidatos = []  # (inicio, fin, texto_nuevo, forzar_negrita)

    for patron, sustituto in cambios.items():
        for m in re.finditer(patron, texto_original, flags=re.IGNORECASE):
            candidatos.append((m.start(), m.end(), sustituto, False))

    for frase in NEG_FRASES:
        for m in re.finditer(re.escape(frase), texto_original):
            candidatos.append((m.start(), m.end(), frase, True))

    for m in re.finditer(r"\[([^\]]+)\]", texto_original):
        campo = m.group(1)
        valor = reemplazos.get(campo, m.group(0))
        forzar = normalizar(campo) in NEG_CAMPOS
        candidatos.append((m.start(), m.end(), str(valor), forzar))

    # Ordenamos por posición de inicio y descartamos solapes (no deberían darse,
    # pero por seguridad nos quedamos con el primero que aparece)
    candidatos.sort(key=lambda c: c[0])
    segmentos = []
    fin_anterior = -1
    for inicio, fin, nuevo, forzar in candidatos:
        if inicio < fin_anterior:
            continue
        segmentos.append((inicio, fin, nuevo, forzar))
        fin_anterior = fin
    return segmentos

def reconstruir_parrafo(parrafo, segmentos):
    if not parrafo.runs: return

    # Guardamos una copia del formato de cada run original
    runs_guardados = []
    for r in parrafo.runs:
        runs_guardados.append({
            "bold": r.bold,
            "italic": r.italic,
            "underline": r.underline,
            "font_name": r.font.name,
            "font_size": r.font.size,
            "font_color": r.font.color.rgb if r.font.color and r.font.color.type else None,
            "highlight": r.font.highlight_color,
            "strike": r.font.strike,
            "subscript": r.font.subscript,
            "superscript": r.font.superscript,
            "texto_orig": r.text,
        })

    texto_original = "".join(rd["texto_orig"] for rd in runs_guardados)
    fmt_base = runs_guardados[0] if runs_guardados else {}

    # Mapa de (inicio, fin, formato) por posición de carácter en el texto ORIGINAL
    bold_map = []
    pos = 0
    for rd in runs_guardados:
        n = len(rd["texto_orig"])
        bold_map.append((pos, pos + n, rd))
        pos += n

    def fmt_para_pos(posicion):
        for inicio, fin, rd in bold_map:
            if inicio <= posicion < fin:
                return rd
        return fmt_base

    def aplicar_formato(run, fmt, forzar_negrita=False):
        run.bold = True if forzar_negrita else fmt.get("bold")
        run.italic = fmt.get("italic")
        run.underline = fmt.get("underline")
        run.font.name = fmt.get("font_name")
        run.font.size = fmt.get("font_size")
        if fmt.get("font_color"):
            run.font.color.rgb = fmt.get("font_color")
        run.font.highlight_color = fmt.get("highlight")
        run.font.strike = fmt.get("strike")
        run.font.subscript = fmt.get("subscript")
        run.font.superscript = fmt.get("superscript")

    def emitir_tramo_original(ini, fin):
        """Copia texto_original[ini:fin] respetando el formato original carácter a carácter."""
        p = ini
        while p < fin:
            fmt = fmt_para_pos(p)
            fin_run = next((f for (i, f, rd) in bold_map if i <= p < f), fin)
            corte = min(fin_run, fin)
            trozo = texto_original[p:corte]
            if trozo:
                run = parrafo.add_run(trozo)
                aplicar_formato(run, fmt, forzar_negrita=False)
            p = corte

    # Guardamos los elementos XML originales para poder eliminarlos al final
    # (en vez de dejarlos vacíos, lo que ensuciaría el documento)
    elementos_originales = [r._element for r in parrafo.runs]

    pos_actual = 0
    for inicio, fin, nuevo, forzar in segmentos:
        emitir_tramo_original(pos_actual, inicio)
        if nuevo:
            fmt = fmt_para_pos(inicio)
            run = parrafo.add_run(nuevo)
            aplicar_formato(run, fmt, forzar_negrita=forzar)
        pos_actual = fin
    emitir_tramo_original(pos_actual, len(texto_original))

    # Eliminamos los runs originales (los nuevos ya se añadieron al final del párrafo)
    for el in elementos_originales:
        el.getparent().remove(el)

def procesar_todo(doc, reemplazos, sexo):
    for p in doc.paragraphs:
        segmentos = calcular_segmentos(p.text, sexo, reemplazos)
        if segmentos:
            reconstruir_parrafo(p, segmentos)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    segmentos = calcular_segmentos(p.text, sexo, reemplazos)
                    if segmentos:
                        reconstruir_parrafo(p, segmentos)

# --- INTERFAZ GRÁFICA ---
class AppCertificadosCentrado:
    def __init__(self, root):
        self.root = root
        self.root.title("Generador de Títulos")
        # 1. Definimos el tamaño que queremos
        ancho = 520
        alto = 400
        
        # 2. Calculamos la posición (Esto es lo que va dentro de init)
        pos_x = (self.root.winfo_screenwidth() // 2) - (ancho // 2)
        
        # 3. Aplicamos la magia: Ancho x Alto + Derecha + Arriba (20)
        self.root.geometry(f"{ancho}x{alto}+{pos_x}+20")
        self.root.resizable(False, False)
        
        self.plantilla = ""
        self.excel = ""

        # Título centrado
        tk.Label(root, text="🎓 Generador de Certificados de Título", 
                 font=("Segoe UI", 14, "bold")).pack(pady=(30, 20))

        # Contenedor central para alineación
        container = tk.Frame(root)
        container.pack(expand=True)

        # Botón 1 y su tick
        f1 = tk.Frame(container)
        f1.pack(pady=5)
        self.btn_p = tk.Button(f1, text="1. Seleccionar Plantilla Word", command=self.sel_p, width=30)
        self.btn_p.pack(side="left")
        self.tick_p = tk.Label(f1, text="", fg="#28a745", font=("Arial", 10, "bold"), width=20, anchor="w")
        self.tick_p.pack(side="left", padx=5)

        # Botón 2 y su tick
        f2 = tk.Frame(container)
        f2.pack(pady=5)
        self.btn_e = tk.Button(f2, text="2. Seleccionar Datos Excel", command=self.sel_e, width=30)
        self.btn_e.pack(side="left")
        self.tick_e = tk.Label(f2, text="", fg="#28a745", font=("Arial", 10, "bold"), width=20, anchor="w")
        self.tick_e.pack(side="left", padx=5)

        # Botón Generar 
        self.btn_go = tk.Button(root, text="GENERAR DOCUMENTACIÓN", command=self.hilo, 
                                bg="#88ffff", fg="white", font=("Segoe UI", 11, "bold"), 
                                width=35, height=2, state="disabled", relief="flat")
        self.btn_go.pack(pady=(30, 10))

        self.lbl_status = tk.Label(root, text="Faltan archivos por seleccionar", fg="#6c757d")
        self.lbl_status.pack(pady=(0, 20))

    def verificar(self):
        if self.plantilla and self.excel:
            self.btn_go.config(state="normal", bg="#28a745", cursor="hand2")
            self.lbl_status.config(text="¡Todo listo para generar!", fg="#28a745")

    def sel_p(self):
        self.plantilla = filedialog.askopenfilename(filetypes=[("Word", "*.docx")])
        if self.plantilla:
            nombre = os.path.basename(self.plantilla)
            self.tick_p.config(text=f"✔️ {nombre[:15]}...")
            self.verificar()

    def sel_e(self):
        self.excel = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx")])
        if self.excel:
            nombre = os.path.basename(self.excel)
            self.tick_e.config(text=f"✔️ {nombre[:15]}...")
            self.verificar()

    def hilo(self):
        dest = filedialog.askdirectory(title="Seleccionar carpeta de guardado")
        if dest:
            self.btn_go.config(state="disabled", text="Procesando...", bg="#6c757d")
            threading.Thread(target=self.run, args=(dest,), daemon=True).start()

    def run(self, folder):
        try:
            df = pd.read_excel(self.excel)
            df.columns = [str(c).strip() for c in df.columns]
            cols_norm = {normalizar(c):c for c in df.columns}
            
            doc_temp = Document(self.plantilla)
            full_text = "".join([p.text for p in doc_temp.paragraphs])
            marcadores = set(re.findall(r"\[([^\]]+)\]", full_text))

            # Clave del marcador "nombre" tal y como aparece en la plantilla
            # (puede estar escrito como [nombre], [Nombre], [NOMBRE]...)
            clave_nombre = next((m for m in marcadores if normalizar(m) == "nombre"), None)

            nombres_usados = {}
            for idx, fila in df.iterrows():
                doc = Document(self.plantilla)
                col_sexo = cols_norm.get("sexo", "")
                sexo = fila.get(col_sexo, "F") if col_sexo else "F"
                if pd.isna(sexo) or not str(sexo).strip():
                    sexo = "F"
                reemplazos = {}
                for m in marcadores:
                    col = cols_norm.get(normalizar(m), "")
                    val = fila.get(col, "") if col else ""
                    if "fecha" in normalizar(m): val = fecha_castellano(val)
                    elif "media" in normalizar(m):
                        try: val = f"{float(val):.2f}"
                        except: val = str(val)
                    reemplazos[m] = "" if pd.isna(val) else str(val)

                procesar_todo(doc, reemplazos, sexo)

                # Nombre base del archivo: el valor ya resuelto del campo [nombre]
                # (si la plantilla lo tiene); si no, recurrimos a la columna "nombre" del Excel
                if clave_nombre and reemplazos.get(clave_nombre):
                    nombre_persona = reemplazos[clave_nombre]
                else:
                    nombre_persona = str(fila.get("nombre", f"Cert_{idx}"))

                nom = "".join(c for c in nombre_persona if c not in r'\/:*?"<>|').strip()
                if not nom:
                    nom = f"Cert_{idx}"

                # Evitamos sobrescribir si hay dos personas con el mismo nombre
                veces = nombres_usados.get(nom, 0)
                nombres_usados[nom] = veces + 1
                nombre_archivo = nom if veces == 0 else f"{nom} ({veces + 1})"

                doc.save(os.path.join(folder, f"{nombre_archivo}.docx"))

            messagebox.showinfo("Hecho", f"Se han generado {len(df)} certificados con éxito.")
            os.startfile(folder)
        except Exception as e:
            messagebox.showerror("Error", f"Ocurrió un error: {e}")
        
        self.btn_go.config(state="normal", text="GENERAR DOCUMENTACIÓN", bg="#28a745")

if __name__ == "__main__":
    root = tk.Tk()
    AppCertificadosCentrado(root)
    root.mainloop()