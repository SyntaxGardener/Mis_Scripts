"""
TXT → DOCX / PDF Formateador  v4.0
Convierte archivos .txt con marcadores de formato a Word y/o PDF.

Novedades v4:
  - Tablas con ancho de columna adaptado al contenido
  - Opción de abrir carpeta de destino y/o el archivo al terminar
  - Detección inteligente: "- A) texto" se trata como lista, no viñeta

Marcadores reconocidos:
  **texto**     → negrita          *texto*   → cursiva
  ==texto==     → resaltado        # / ## / ### → encabezados
  ▸ • ● - *    → viñetas          a) 1. 2.  → listas numeradas
  ---           → separador        | col |   → tabla Markdown

Requisitos: pip install python-docx reportlab
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, colorchooser
import re, os, subprocess, sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    HRFlowable, Table, TableStyle)
    HAS_PDF = True
except ImportError:
    HAS_PDF = False


# ═══════════════════════════════════════════════════════════════════════════════
#  PARSER
# ═══════════════════════════════════════════════════════════════════════════════

# Patrón que indica que una línea es "- A) texto" o "- 1. texto" etc.
# → se debe tratar como lista, suprimiendo la viñeta
_RE_BULLET_LIST = re.compile(
    r'^(\s*)[\-\*]\s+([A-Za-z]\)|[IVXLC]+[\.\)]|\d+[\.\)])\s+(.*)'
)

def clasificar_linea(linea):
    raw = linea.rstrip('\n')
    if raw.strip() == '':
        return ('empty', '', 0)
    if re.match(r'^[\-=]{3,}\s*$', raw.strip()):
        return ('separator', '', 0)

    # Encabezado Markdown
    m = re.match(r'^(#{1,3})\s+(.*)', raw)
    if m:
        return ('heading', m.group(2).strip(), len(m.group(1)))

    # [TÍTULO EN CORCHETES]
    m = re.match(r'^\[([A-ZÁÉÍÓÚÑ][^\]]+)\]\s*$', raw.strip())
    if m:
        return ('heading', m.group(1).strip(), 2)

    # ==TÍTULO EN MAYÚSCULAS==
    s = raw.strip()
    if s.startswith('==') and s.endswith('==') and len(s) > 6:
        inner = s[2:-2].strip()
        if inner.isupper() or re.match(r'^[A-ZÁÉÍÓÚÑ\s\d\:\-]+$', inner):
            return ('heading', inner, 2)

    # "- A) texto" / "- 1. texto"  → lista (suprime viñeta)
    m = _RE_BULLET_LIST.match(raw)
    if m:
        indent  = len(m.group(1)) // 2
        marker  = m.group(2)          # "A)", "1.", "B)", …
        content = m.group(3).strip()
        # Detectar si es letra (A-D) = opción múltiple → numbered con letra
        return ('numbered', f"{marker} {content}", indent)

    # Viñeta con símbolo especial
    bp = re.match(r'^(\s*)(▸|•|●|○|■|□|►|☆|✓|✗|❌|💡|⚠|–|—)\s+(.*)', raw)
    if bp:
        return ('bullet', bp.group(3).strip(), len(bp.group(1)) // 2)

    # Viñeta con - o *  (solo si no es separador ni bullet+lista)
    m = re.match(r'^(\s*)[\-\*]\s+((?!\-\-).+)', raw)
    if m:
        return ('bullet', m.group(2).strip(), len(m.group(1)) // 2)

    # Lista numerada o alfabética directa: "a)", "1.", "A)"
    m = re.match(r'^(\s*)([A-Za-z]\)|\d+[\.\)])\s+(.*)', raw)
    if m:
        marker  = m.group(2)
        content = m.group(3).strip()
        return ('numbered', f"{marker} {content}", len(m.group(1)) // 2)

    return ('normal', raw, 0)


def parsear_inline(texto):
    segs, patron, pos = [], re.compile(r'==(.+?)==|\*\*(.+?)\*\*|\*(.+?)\*'), 0
    for m in patron.finditer(texto):
        if m.start() > pos:
            segs.append((texto[pos:m.start()], False, False, False))
        if   m.group(1): segs.append((m.group(1), False, False, True))
        elif m.group(2): segs.append((m.group(2), True,  False, False))
        elif m.group(3): segs.append((m.group(3), False, True,  False))
        pos = m.end()
    if pos < len(texto):
        segs.append((texto[pos:], False, False, False))
    return segs or [(texto, False, False, False)]


def hex_to_rgb(h):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


# ─── Tablas ──────────────────────────────────────────────────────────────────

def es_linea_tabla(l):
    return '|' in l.strip()

def es_separador_tabla(l):
    s = l.strip().replace(' ', '')
    return bool(re.match(r'^[\|\-\:]+$', s)) and '-' in s and len(s) > 2

def parsear_fila(l):
    s = l.strip()
    if s.startswith('|'): s = s[1:]
    if s.endswith('|'):   s = s[:-1]
    return [c.strip() for c in s.split('|')]

def limpiar_marcadores(texto):
    """Quita marcadores inline para medir longitud de texto plano."""
    texto = re.sub(r'==(.+?)==',      r'\1', texto)
    texto = re.sub(r'\*\*(.+?)\*\*', r'\1', texto)
    texto = re.sub(r'\*(.+?)\*',     r'\1', texto)
    return texto

def calcular_anchos_columnas(filas, ancho_total, min_pct=0.10):
    """
    Calcula anchos proporcionales al contenido de cada columna.
    min_pct: fracción mínima por columna (evita columnas microscópicas).
    """
    if not filas:
        return []
    ncols = max(len(f) for f in filas)
    # Longitud máxima de texto en cada columna
    maxlen = [0] * ncols
    for fila in filas:
        for ci, celda in enumerate(fila):
            l = len(limpiar_marcadores(celda))
            if l > maxlen[ci]:
                maxlen[ci] = l
    total = sum(maxlen) or 1
    # Porcentajes con mínimo garantizado
    pcts = [max(min_pct, ml / total) for ml in maxlen]
    # Renormalizar a 1.0
    suma = sum(pcts)
    pcts = [p / suma for p in pcts]
    return [ancho_total * p for p in pcts]


def agrupar_bloques(lineas):
    """
    Agrupa líneas en bloques clasificados.
    Las tablas se devuelven como ('table', filas).
    """
    bloques = []
    i = 0
    while i < len(lineas):
        linea = lineas[i]
        if es_linea_tabla(linea) and not es_separador_tabla(linea):
            filas_raw = []
            j = i
            while j < len(lineas) and es_linea_tabla(lineas[j]):
                filas_raw.append(lineas[j])
                j += 1
            filas = [parsear_fila(f) for f in filas_raw
                     if not es_separador_tabla(f)]
            if filas:
                bloques.append(('table', filas))
            i = j
        else:
            bloques.append(clasificar_linea(linea))
            i += 1
    return bloques


# ═══════════════════════════════════════════════════════════════════════════════
#  GENERADOR DOCX
# ═══════════════════════════════════════════════════════════════════════════════

def agregar_highlight(run):
    rPr = run._r.get_or_add_rPr()
    hl  = OxmlElement('w:highlight')
    hl.set(qn('w:val'), 'yellow')
    rPr.append(hl)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color.lstrip('#'))
    tcPr.append(shd)

def set_cell_width(cell, width_dxa):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(width_dxa)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def generar_docx(lineas, ruta_salida, fuente, colores_h):
    if not HAS_DOCX:
        raise ImportError("Instala python-docx: pip install python-docx")

    doc = Document()
    doc.styles['Normal'].font.name = fuente
    doc.styles['Normal'].font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.5)
        sec.left_margin  = Cm(3)
        sec.right_margin = Cm(2.5)

    # Ancho útil en DXA (1 DXA = 1/1440 pulgada)
    # A4 = 21 cm, márgenes 3+2.5 cm → 15.5 cm → en DXA:
    CONTENT_DXA = int(15.5 / 2.54 * 1440)   # ≈ 8773

    def inline(p, texto, hdr=False):
        for (t, bold, italic, hl) in parsear_inline(texto):
            r = p.add_run(t)
            r.bold      = bold or hdr
            r.italic    = italic
            r.font.name = fuente
            if hl:
                agregar_highlight(r)

    def add_tabla_docx(filas):
        if not filas:
            return
        ncols   = max(len(f) for f in filas)
        # Anchos proporcionales al contenido, en DXA
        anchos  = calcular_anchos_columnas(filas, CONTENT_DXA)
        # Rellenar si hay menos anchos que columnas (por seguridad)
        while len(anchos) < ncols:
            anchos.append(CONTENT_DXA / ncols)

        tbl = doc.add_table(rows=len(filas), cols=ncols)
        tbl.style = 'Table Grid'

        for ri, fila in enumerate(filas):
            hdr = (ri == 0)
            for ci in range(ncols):
                cell  = tbl.cell(ri, ci)
                set_cell_width(cell, anchos[ci])
                texto = fila[ci] if ci < len(fila) else ''
                if hdr:
                    set_cell_bg(cell, '#D9E1F2')
                p = cell.paragraphs[0]
                p.clear()
                for (t, bold, italic, hl2) in parsear_inline(texto):
                    r = p.add_run(t)
                    r.bold      = bold or hdr
                    r.italic    = italic
                    r.font.name = fuente
                    r.font.size = Pt(10)
                    if hl2:
                        agregar_highlight(r)
        doc.add_paragraph()

    bloques = agrupar_bloques(lineas)
    for b in bloques:
        tipo = b[0]

        if tipo == 'table':
            add_tabla_docx(b[1])

        elif tipo == 'empty':
            continue

        elif tipo == 'separator':
            p   = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bot  = OxmlElement('w:bottom')
            bot.set(qn('w:val'),   'single')
            bot.set(qn('w:sz'),    '6')
            bot.set(qn('w:space'), '1')
            bot.set(qn('w:color'), 'AAAAAA')
            pBdr.append(bot)
            pPr.append(pBdr)

        elif tipo == 'heading':
            _, contenido, nivel = b
            nw  = min(nivel, 3)
            p   = doc.add_heading(level=nw)
            p.clear()
            run = p.add_run(contenido)
            run.bold      = True
            run.font.name = fuente
            run.font.size = Pt([0, 18, 14, 12][nw])
            r2, g2, b2    = hex_to_rgb(colores_h.get(nw, '#1F3864'))
            run.font.color.rgb = RGBColor(r2, g2, b2)
            p.paragraph_format.space_before = Pt(12 if nivel == 1 else 8)
            p.paragraph_format.space_after  = Pt(4)

        elif tipo == 'bullet':
            _, contenido, nivel = b
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (nivel + 1))
            inline(p, contenido)

        elif tipo == 'numbered':
            _, contenido, nivel = b
            # Usar párrafo normal con sangría para preservar el marcador
            # (A), B), 1., 2. ya vienen en el contenido)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent   = Inches(0.35 * (nivel + 1))
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after   = Pt(3)
            inline(p, contenido)

        else:  # normal
            _, contenido, _ = b
            p = doc.add_paragraph()
            inline(p, contenido)
            p.paragraph_format.space_after = Pt(4)

    doc.save(ruta_salida)


# ═══════════════════════════════════════════════════════════════════════════════
#  GENERADOR PDF
# ═══════════════════════════════════════════════════════════════════════════════

def inline_a_html(texto):
    texto = texto.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    texto = re.sub(r'==(.+?)==',      r'<font backColor="#FFFF00"><b>\1</b></font>', texto)
    texto = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', texto)
    texto = re.sub(r'\*(.+?)\*',     r'<i>\1</i>', texto)
    return texto


def generar_pdf(lineas, ruta_salida, fuente_pdf, colores_h):
    if not HAS_PDF:
        raise ImportError("Instala reportlab: pip install reportlab")

    mapa = {'Calibri':'Helvetica','Arial':'Helvetica','Times New Roman':'Times-Roman',
            'Georgia':'Times-Roman','Garamond':'Times-Roman','Courier New':'Courier',
            'Verdana':'Helvetica','Trebuchet MS':'Helvetica'}
    fn = mapa.get(fuente_pdf, 'Helvetica')
    fb = {'Helvetica':'Helvetica-Bold','Times-Roman':'Times-Bold',
          'Courier':'Courier-Bold'}.get(fn, 'Helvetica-Bold')

    def rl_col(hx):
        r, g, b = hex_to_rgb(hx)
        return colors.Color(r/255, g/255, b/255)

    PAGE_W = A4[0] - 3*cm - 2.5*cm

    doc = SimpleDocTemplate(ruta_salida, pagesize=A4,
                             topMargin=2.5*cm, bottomMargin=2.5*cm,
                             leftMargin=3*cm,  rightMargin=2.5*cm)

    nst = ParagraphStyle('N',  fontName=fn, fontSize=10.5, leading=15, spaceAfter=4)
    hst = {
        1: ParagraphStyle('H1', fontName=fb, fontSize=18, spaceAfter=8,
                           spaceBefore=14, textColor=rl_col(colores_h.get(1,'#1F3864'))),
        2: ParagraphStyle('H2', fontName=fb, fontSize=14, spaceAfter=6,
                           spaceBefore=10, textColor=rl_col(colores_h.get(2,'#2E5A8E'))),
        3: ParagraphStyle('H3', fontName=fb, fontSize=12, spaceAfter=4,
                           spaceBefore=8,  textColor=rl_col(colores_h.get(3,'#4472C4'))),
    }
    bst = ParagraphStyle('B',  parent=nst, leftIndent=18, spaceAfter=3)
    # Listas con marcador (A), 1., etc.) — sangría colgante
    lst = ParagraphStyle('L',  parent=nst, leftIndent=28, firstLineIndent=-18,
                          spaceAfter=3)
    cn  = ParagraphStyle('TC', fontName=fn, fontSize=9, leading=12)
    ch  = ParagraphStyle('TH', fontName=fb, fontSize=9, leading=12)

    def hacer_tabla_pdf(filas):
        if not filas:
            return None
        ncols   = max(len(f) for f in filas)
        anchos  = calcular_anchos_columnas(filas, PAGE_W)
        while len(anchos) < ncols:
            anchos.append(PAGE_W / ncols)

        data = []
        for ri, fila in enumerate(filas):
            row = []
            for ci in range(ncols):
                txt = fila[ci] if ci < len(fila) else ''
                st  = ch if ri == 0 else cn
                row.append(Paragraph(inline_a_html(txt), st))
            data.append(row)

        t = Table(data, colWidths=anchos, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND',   (0,0), (-1,0),  colors.HexColor('#D9E1F2')),
            ('TEXTCOLOR',    (0,0), (-1,0),  colors.HexColor('#1F3864')),
            ('GRID',         (0,0), (-1,-1), 0.5, colors.HexColor('#AAAAAA')),
            ('LINEBELOW',    (0,0), (-1,0),  1.0, colors.HexColor('#2E5A8E')),
            ('TOPPADDING',   (0,0), (-1,-1), 4),
            ('BOTTOMPADDING',(0,0), (-1,-1), 4),
            ('LEFTPADDING',  (0,0), (-1,-1), 6),
            ('RIGHTPADDING', (0,0), (-1,-1), 6),
            ('ROWBACKGROUNDS',(0,1),(-1,-1),
             [colors.white, colors.HexColor('#F5F8FF')]),
            ('VALIGN',       (0,0), (-1,-1), 'MIDDLE'),
        ]))
        return t

    story   = []
    bloques = agrupar_bloques(lineas)

    for b in bloques:
        tipo = b[0]

        if tipo == 'table':
            t = hacer_tabla_pdf(b[1])
            if t:
                story.append(Spacer(1, 6))
                story.append(t)
                story.append(Spacer(1, 10))

        elif tipo == 'empty':
            story.append(Spacer(1, 4))

        elif tipo == 'separator':
            story += [Spacer(1,4),
                      HRFlowable(width='100%', thickness=0.5,
                                 color=colors.HexColor('#AAAAAA')),
                      Spacer(1,4)]

        elif tipo == 'heading':
            _, contenido, nivel = b
            story.append(Paragraph(inline_a_html(contenido), hst[min(nivel,3)]))

        elif tipo == 'bullet':
            _, contenido, nivel = b
            st = ParagraphStyle(f'bl{nivel}', parent=bst, leftIndent=18+12*nivel)
            story.append(Paragraph(f'• {inline_a_html(contenido)}', st))

        elif tipo == 'numbered':
            _, contenido, nivel = b
            st = ParagraphStyle(f'ls{nivel}', parent=lst,
                                 leftIndent=28+12*nivel)
            story.append(Paragraph(inline_a_html(contenido), st))

        else:
            _, contenido, _ = b
            story.append(Paragraph(inline_a_html(contenido), nst))

    doc.build(story)


# ═══════════════════════════════════════════════════════════════════════════════
#  GUI
# ═══════════════════════════════════════════════════════════════════════════════

FUENTES = ['Arial', 'Calibri', 'Courier New', 'Garamond',
           'Georgia', 'Times New Roman', 'Trebuchet MS', 'Verdana']

COLORES_DEFECTO = {1: '#1F3864', 2: '#2E5A8E', 3: '#4472C4'}
LABEL_NIVEL     = {1: 'H1 (principal)', 2: 'H2 (sección)', 3: 'H3 (subsección)'}


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("✦  TXT → DOCX / PDF  |  Formateador de documentos")
        self.resizable(False, False)
        self.colores_h = dict(COLORES_DEFECTO)
        self._construir_ui()
        self._comprobar_dependencias()
        self.update_idletasks()
        w, h = 680, 600
        x = (self.winfo_screenwidth() - w) // 2
        self.geometry(f"{w}x{h}+{x}+5")

    # ── Construcción UI ───────────────────────────────────────────────────────

    def _construir_ui(self):
        P = dict(padx=12, pady=5)

        # Cabecera
        cab = tk.Frame(self, bg='#1F3864')
        cab.pack(fill='x')
        tk.Label(cab, text="✦  Formateador de documentos  ·  TXT → DOCX / PDF  v4",
                 bg='#1F3864', fg='white', font=('Arial', 13, 'bold'), pady=10).pack()

        # Entrada
        fe = ttk.LabelFrame(self, text="Archivo .txt de entrada")
        fe.pack(fill='x', **P)
        self.var_entrada = tk.StringVar()
        ttk.Entry(fe, textvariable=self.var_entrada, width=72).pack(side='left', padx=6, pady=6)
        ttk.Button(fe, text="…", width=3, command=self._sel_entrada).pack(side='left')

        # Salida
        fs = ttk.LabelFrame(self, text="Carpeta de salida")
        fs.pack(fill='x', **P)
        self.var_salida = tk.StringVar()
        ttk.Entry(fs, textvariable=self.var_salida, width=72).pack(side='left', padx=6, pady=6)
        ttk.Button(fs, text="…", width=3, command=self._sel_salida).pack(side='left')

        # Opciones de salida
        fo = ttk.LabelFrame(self, text="Opciones de salida")
        fo.pack(fill='x', **P)

        # Fila 1: formato
        f1 = ttk.Frame(fo)
        f1.pack(fill='x', padx=8, pady=(6,2))
        ttk.Label(f1, text="Formato:").pack(side='left')
        self.var_docx = tk.BooleanVar(value=True)
        self.var_pdf  = tk.BooleanVar(value=False)
        self.cb_docx  = ttk.Checkbutton(f1, text="Word (.docx)", variable=self.var_docx)
        self.cb_docx.pack(side='left', padx=(8,16))
        self.cb_pdf   = ttk.Checkbutton(f1, text="PDF", variable=self.var_pdf)
        self.cb_pdf.pack(side='left')

        # Fila 2: fuente
        f2 = ttk.Frame(fo)
        f2.pack(fill='x', padx=8, pady=(2,4))
        ttk.Label(f2, text="Tipo de letra:").pack(side='left')
        self.var_fuente = tk.StringVar(value='Calibri')
        ttk.Combobox(f2, textvariable=self.var_fuente,
                     values=FUENTES, state='readonly', width=22).pack(side='left', padx=8)

        # Fila 3: al terminar
        f3 = ttk.Frame(fo)
        f3.pack(fill='x', padx=8, pady=(2,8))
        ttk.Label(f3, text="Al terminar:").pack(side='left')
        self.var_abrir_arch    = tk.BooleanVar(value=False)
        self.var_abrir_carpeta = tk.BooleanVar(value=False)
        ttk.Checkbutton(f3, text="Abrir archivo(s)", variable=self.var_abrir_arch).pack(side='left', padx=(8,16))
        ttk.Checkbutton(f3, text="Abrir carpeta de destino", variable=self.var_abrir_carpeta).pack(side='left')

        # Colores encabezados
        fc = ttk.LabelFrame(self, text="Color de encabezados")
        fc.pack(fill='x', **P)
        self._btns_color = {}
        fila_c = ttk.Frame(fc)
        fila_c.pack(padx=8, pady=8)
        for n in (1, 2, 3):
            ttk.Label(fila_c, text=LABEL_NIVEL[n]).grid(row=0, column=(n-1)*2, padx=(10,4))
            btn = tk.Button(fila_c, width=5, bg=self.colores_h[n],
                            relief='solid', bd=1, cursor='hand2',
                            command=lambda nivel=n: self._elegir_color(nivel))
            btn.grid(row=0, column=(n-1)*2+1, padx=(0,18))
            self._btns_color[n] = btn

        # Vista previa
        fv = ttk.LabelFrame(self, text="Vista previa (estructura detectada)")
        fv.pack(fill='both', expand=True, **P)
        self.texto_prev = tk.Text(fv, height=9, wrap='word', font=('Consolas', 9),
                                   bg='#F8F8F8', relief='flat', state='disabled')
        sc = ttk.Scrollbar(fv, command=self.texto_prev.yview)
        self.texto_prev.configure(yscrollcommand=sc.set)
        self.texto_prev.pack(side='left', fill='both', expand=True, padx=4, pady=4)
        sc.pack(side='right', fill='y', pady=4)

        # Botones
        fb2 = ttk.Frame(self)
        fb2.pack(fill='x', padx=12, pady=6)
        ttk.Button(fb2, text="Vista previa", command=self._previsualizar).pack(side='left', padx=4)
        ttk.Button(fb2, text="✦  Convertir", command=self._convertir).pack(side='right', padx=4)

        self.var_estado = tk.StringVar(value="Listo.")
        ttk.Label(self, textvariable=self.var_estado, foreground='#555').pack(pady=(0,6))

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _comprobar_dependencias(self):
        avisos = []
        if not HAS_DOCX:
            avisos.append("python-docx no instalado  →  pip install python-docx")
            self.cb_docx.config(state='disabled'); self.var_docx.set(False)
        if not HAS_PDF:
            avisos.append("reportlab no instalado  →  pip install reportlab")
            self.cb_pdf.config(state='disabled');  self.var_pdf.set(False)
        if avisos:
            messagebox.showwarning("Dependencias faltantes", "\n".join(avisos))

    def _elegir_color(self, nivel):
        res = colorchooser.askcolor(color=self.colores_h[nivel],
                                    title=f"Color para {LABEL_NIVEL[nivel]}")
        if res and res[1]:
            self.colores_h[nivel] = res[1]
            self._btns_color[nivel].config(bg=res[1])

    def _sel_entrada(self):
        ruta = filedialog.askopenfilename(
            title="Seleccionar archivo .txt",
            filetypes=[("Archivos de texto", "*.txt"), ("Todos", "*.*")])
        if ruta:
            self.var_entrada.set(ruta)
            if not self.var_salida.get():
                self.var_salida.set(os.path.dirname(ruta))
            self._previsualizar()

    def _sel_salida(self):
        ruta = filedialog.askdirectory(title="Seleccionar carpeta de salida")
        if ruta:
            self.var_salida.set(ruta)

    def _leer_lineas(self):
        ruta = self.var_entrada.get()
        if not ruta or not os.path.isfile(ruta):
            return None
        with open(ruta, encoding='utf-8', errors='replace') as f:
            return f.readlines()

    def _previsualizar(self):
        lineas = self._leer_lineas()
        if not lineas:
            return
        self.texto_prev.config(state='normal')
        self.texto_prev.delete('1.0', 'end')
        for tag, col in [('heading','#1F3864'), ('bullet','#2E7D32'),
                          ('numbered','#6A1B9A'), ('separator','#888'),
                          ('table','#8B4513'),   ('normal','#333')]:
            self.texto_prev.tag_configure(tag, foreground=col)
        iconos = {'bullet':'  •', 'numbered':'  №', 'separator':'────',
                  'table':'  ▦', 'normal':'  ¶'}
        bloques = agrupar_bloques(lineas)
        for b in bloques:
            tipo = b[0]
            if tipo == 'empty':
                self.texto_prev.insert('end', '\n'); continue
            if tipo == 'table':
                filas = b[1]
                r = f"  ▦  TABLA  {len(filas)} filas × {max(len(f) for f in filas)} columnas"
                self.texto_prev.insert('end', r + '\n', 'table'); continue
            _, contenido, nivel = b
            prefijo = f'▶ H{nivel}' if tipo == 'heading' else iconos.get(tipo, '  ')
            self.texto_prev.insert('end', f"{prefijo:6}  {contenido[:90]}\n", tipo)
        self.texto_prev.config(state='disabled')
        n_tab = sum(1 for b in bloques if b[0]=='table')
        self.var_estado.set(
            f"Vista previa: {len(lineas)} líneas — {n_tab} tabla(s) detectada(s).")

    def _abrir_ruta(self, ruta):
        try:
            if sys.platform == 'win32':    os.startfile(ruta)
            elif sys.platform == 'darwin': subprocess.call(['open', ruta])
            else:                          subprocess.call(['xdg-open', ruta])
        except Exception:
            pass

    def _convertir(self):
        txt     = self.var_entrada.get()
        carpeta = self.var_salida.get()
        fuente  = self.var_fuente.get()

        if not txt or not os.path.isfile(txt):
            messagebox.showerror("Error", "Selecciona un archivo .txt válido."); return
        if not carpeta or not os.path.isdir(carpeta):
            messagebox.showerror("Error", "Selecciona una carpeta de salida válida."); return
        if not self.var_docx.get() and not self.var_pdf.get():
            messagebox.showerror("Error", "Marca al menos un formato de salida."); return

        lineas = self._leer_lineas()
        base   = os.path.splitext(os.path.basename(txt))[0]
        generados, errores = [], []

        if self.var_docx.get():
            ruta = os.path.join(carpeta, f"{base}_formateado.docx")
            try:    generar_docx(lineas, ruta, fuente, self.colores_h); generados.append(ruta)
            except Exception as e: errores.append(f"DOCX: {e}")

        if self.var_pdf.get():
            ruta = os.path.join(carpeta, f"{base}_formateado.pdf")
            try:    generar_pdf(lineas, ruta, fuente, self.colores_h); generados.append(ruta)
            except Exception as e: errores.append(f"PDF: {e}")

        if generados:
            msg = "Archivos generados:\n" + "\n".join(f"  ✓ {r}" for r in generados)
            if errores: msg += "\n\nErrores:\n" + "\n".join(errores)
            messagebox.showinfo("¡Listo!", msg)
            self.var_estado.set("Convertido: " + ", ".join(os.path.basename(r) for r in generados))
            # Abrir archivos
            if self.var_abrir_arch.get():
                for r in generados:
                    self._abrir_ruta(r)
            # Abrir carpeta
            if self.var_abrir_carpeta.get():
                self._abrir_ruta(carpeta)
        else:
            messagebox.showerror("Error", "\n".join(errores))
            self.var_estado.set("Error en la conversión.")


if __name__ == '__main__':
    App().mainloop()
