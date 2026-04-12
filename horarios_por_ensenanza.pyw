import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
from docx import Document
from docx.shared import RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import unicodedata
from collections import defaultdict
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False


# ── Utilidades ────────────────────────────────────────────────────────────────

def eliminar_tildes(cadena):
    if not cadena:
        return ""
    s = "".join(
        c for c in unicodedata.normalize("NFD", str(cadena))
        if unicodedata.category(c) != "Mn"
    )
    return s.upper().strip()


def aplicar_color_espad(run, texto):
    t = eliminar_tildes(texto)
    if "1.2" in t:
        run.font.color.rgb = RGBColor(255, 0, 0)
    elif "2.1" in t:
        run.font.color.rgb = RGBColor(0, 128, 0)
    elif "2.2" in t:
        run.font.color.rgb = RGBColor(0, 0, 255)
    else:
        run.font.color.rgb = RGBColor(128, 0, 128)


def extraer_ensenanzas_y_profesores(ruta_docx):
    doc = Document(ruta_docx)
    patron = re.compile(
        r"([A-ZÁÉÍÓÚÜÑ][A-ZÁÉÍÓÚÜÑ0-9 \.\-]+?)\s*\(([^)]+)\)",
        re.IGNORECASE | re.UNICODE,
    )
    resultado = defaultdict(set)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for c_idx, celda in enumerate(fila.cells):
                if c_idx == 0:
                    continue
                for linea in celda.text.splitlines():
                    linea = re.sub(r"\b\d{1,2}:\d{2}\b", "", linea)
                    linea = re.sub(r"\bhasta\b", "", linea, flags=re.IGNORECASE)
                    linea = linea.strip()
                    if not linea:
                        continue
                    for m in patron.finditer(linea):
                        mat = m.group(1).strip().rstrip(".,- ")
                        prf = m.group(2).strip()
                        if len(mat) < 2 or mat.isdigit():
                            continue
                        resultado[eliminar_tildes(mat)].add(eliminar_tildes(prf))

    return dict(resultado)



def pdf_a_docx(ruta_pdf):
    """
    Convierte el PDF a un Document de python-docx en memoria,
    replicando la estructura de tabla para que el resto del script
    funcione igual que con un Word original.
    Devuelve la ruta de un docx temporal.
    """
    import tempfile, os
    from docx import Document as DocxDoc
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    with pdfplumber.open(ruta_pdf) as pdf:
        page = pdf.pages[0]
        tables = page.extract_tables()
        if not tables:
            raise ValueError("No se encontró ninguna tabla en el PDF.")
        raw = tables[0]

    # Colapsar columnas duplicadas (celdas fusionadas → grupos de 3)
    n_cols = len(raw[0])
    step = 3
    col_indices = list(range(0, n_cols, step))
    collapsed = []
    for row in raw:
        new_row = []
        for ci in col_indices:
            val = None
            for offset in range(step):
                if ci + offset < n_cols and row[ci + offset]:
                    val = row[ci + offset]
                    break
            new_row.append(val or "")
        collapsed.append(new_row)

    # Eliminar fila separadora vacía entre mañana y tarde
    collapsed = [r for r in collapsed if any(c.strip() for c in r)]

    doc = DocxDoc()
    n_filas = len(collapsed)
    n_cols2 = len(collapsed[0])
    tabla = doc.add_table(rows=n_filas, cols=n_cols2)
    tabla.style = "Table Grid"

    for r_idx, row in enumerate(collapsed):
        for c_idx, val in enumerate(row):
            cell = tabla.cell(r_idx, c_idx)
            # Reemplazar saltos de línea internos por párrafos separados
            partes = val.split("\n") if val else [""]
            cell.paragraphs[0].text = partes[0]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for parte in partes[1:]:
                p = cell.add_paragraph(parte)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    doc.save(tmp.name)
    return tmp.name


# ── Aplicación ────────────────────────────────────────────────────────────────

class AppAuto:
    def __init__(self, root):
        self.root = root
        self.root.title("Generador de Horarios – Detección Automática")
        ancho, alto = 1060, 640
        px = (root.winfo_screenwidth() // 2) - (ancho // 2)
        root.geometry(f"{ancho}x{alto}+{px}+20")

        self.ruta_docx = tk.StringVar()
        self.abrir_carpeta = tk.BooleanVar(value=True)
        self._pdf_tmp = None
        self.ruta_display = ""
        self.grupos = {}      # nombre_grupo -> {"materias": set, "profes": set}
        self.vars_sel = {}    # nombre_grupo -> BooleanVar
        self.filas_ui = {}    # nombre_grupo -> dict de widgets

        # ── Cabecera ──────────────────────────────────────────────────────────
        tk.Label(root, text="Generador de Horarios por Enseñanzas",
                 font=("Arial", 13, "bold")).pack(pady=8)

        # ── 1. Archivo ────────────────────────────────────────────────────────
        frm_file = tk.LabelFrame(root, text=" 1. Selecciona el archivo (Word o PDF) ",
                                 padx=10, pady=6)
        frm_file.pack(fill=tk.X, padx=20, pady=4)
        self.ruta_display_var = tk.StringVar()
        tk.Entry(frm_file, textvariable=self.ruta_display_var, width=65,
                 state="readonly").pack(side=tk.LEFT, expand=True, fill=tk.X)
        tk.Button(frm_file, text="Examinar…", command=self.elegir_archivo,
                  bg="#0078D7", fg="white").pack(side=tk.LEFT, padx=8)
        self.lbl_pdf_aviso = tk.Label(
            root,
            text="⚠  Archivo PDF: los horarios generados saldrán en blanco y negro, sin formato de color.",
            fg="#B8860B", font=("Arial", 9, "italic"), anchor="w")
        # se muestra solo cuando se carga un PDF (ver elegir_archivo)

        # ── Ayuda ────────────────────────────────────────────────────────────
        frm_ayuda = tk.LabelFrame(root, text=" ℹ  Cómo usar los botones ",
                                  padx=10, pady=5)
        frm_ayuda.pack(fill=tk.X, padx=20, pady=2)
        ayuda_txt = (
            "🔗 Fusionar: une varias enseñanzas en un solo horario (p.ej. ESPAD 1.2 COM. + ESPAD 1.2 C-T + ESPAD 1.2 SOCIAL  →  ESPAD 1.2).\n"
            "✂  Separar: deshace una fusión anterior, devolviendo las asignaturas por separado.\n"
            "📋 Copiar a grupos: para asignaturas compartidas entre varios horarios (p.ej. ESPAD ASTURIANO):\n"
            "       ⚠  NO la fusiones. Márcala sola y usa 'Copiar a grupos' para añadirla a cada horario que la necesite."
        )
        tk.Label(frm_ayuda, text=ayuda_txt, justify="left",
                 font=("Arial", 8), fg="#333333", wraplength=820).pack(anchor="w")

        # ── 2. Tabla ──────────────────────────────────────────────────────────
        frm_tabla = tk.LabelFrame(
            root,
            text=" 2. Enseñanzas detectadas  "
                 "— marca varias y pulsa 'Fusionar' para agruparlas ",
            padx=8, pady=6)
        frm_tabla.pack(fill=tk.BOTH, expand=True, padx=20, pady=4)

        cab = tk.Frame(frm_tabla)
        cab.pack(fill=tk.X)
        for col, txt, w in [(0, "✔", 3), (1, "Nombre del grupo", 26),
                             (2, "Asignaturas incluidas", 28),
                             (3, "Profesores", 22)]:
            tk.Label(cab, text=txt, width=w, font=("Arial", 9, "bold"),
                     anchor="w").grid(row=0, column=col, sticky="w", padx=2)
        ttk.Separator(frm_tabla, orient="horizontal").pack(fill=tk.X, pady=2)

        self.canvas = tk.Canvas(frm_tabla, highlightthickness=0)
        sb = tk.Scrollbar(frm_tabla, orient="vertical", command=self.canvas.yview)
        self.frm_items = tk.Frame(self.canvas)
        self.frm_items.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.create_window((0, 0), window=self.frm_items, anchor="nw")
        self.canvas.configure(yscrollcommand=sb.set)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb.pack(side=tk.RIGHT, fill=tk.Y)

        tk.Label(self.frm_items,
                 text="(Selecciona un archivo para detectar las enseñanzas)",
                 fg="gray").grid(row=0, column=0, columnspan=4, pady=20)

        # ── 3. Botones ────────────────────────────────────────────────────────
        frm_bot = tk.Frame(root)
        frm_bot.pack(pady=10)

        tk.Button(frm_bot, text="Seleccionar todo",
                  command=self.sel_todo).pack(side=tk.LEFT, padx=4)
        tk.Button(frm_bot, text="Deseleccionar todo",
                  command=self.desel_todo).pack(side=tk.LEFT, padx=4)
        tk.Button(frm_bot, text="🔗  Fusionar seleccionadas",
                  command=self.fusionar,
                  bg="#7B1FA2", fg="white").pack(side=tk.LEFT, padx=10)
        tk.Button(frm_bot, text="✂  Separar grupo",
                  command=self.separar,
                  bg="#E65100", fg="white").pack(side=tk.LEFT, padx=4)
        tk.Button(frm_bot, text="📋  Copiar a grupos",
                  command=self.copiar_a_grupos,
                  bg="#1565C0", fg="white").pack(side=tk.LEFT, padx=4)
        tk.Checkbutton(frm_bot, text="Abrir carpeta al terminar",
                  variable=self.abrir_carpeta).pack(side=tk.LEFT, padx=6)
        tk.Button(frm_bot, text="⚙  GENERAR HORARIOS",
                  command=self.generar,
                  bg="#2E7D32", fg="white",
                  font=("Arial", 10, "bold")).pack(side=tk.LEFT, padx=14)

    # ── Carga ─────────────────────────────────────────────────────────────────

    def elegir_archivo(self):
        path = filedialog.askopenfilename(
            title="Selecciona el Word maestro",
            filetypes=[("Word y PDF", "*.docx *.pdf"), ("Word", "*.docx"), ("PDF", "*.pdf")])
        if not path:
            return
        if path.lower().endswith(".pdf"):
            if not PDF_SUPPORT:
                messagebox.showerror(
                    "pdfplumber no instalado",
                    "Para abrir PDFs instala pdfplumber:\n  pip install pdfplumber")
                return
            try:
                docx_tmp = pdf_a_docx(path)
                self._pdf_tmp = docx_tmp   # guardar para limpieza
                self.ruta_docx.set(docx_tmp)
                self.ruta_display_var.set(path)
                self.lbl_pdf_aviso.pack(fill=tk.X)
            except Exception as e:
                messagebox.showerror("Error al convertir PDF", str(e))
                return
        else:
            self.ruta_docx.set(path)
            self.ruta_display_var.set(path)
            self.lbl_pdf_aviso.pack_forget()
        self.analizar_documento(self.ruta_docx.get())

    def analizar_documento(self, path):
        for w in self.frm_items.winfo_children():
            w.destroy()
        self.grupos.clear()
        self.vars_sel.clear()
        self.filas_ui.clear()

        try:
            detectado = extraer_ensenanzas_y_profesores(path)
        except Exception as e:
            messagebox.showerror("Error al leer el archivo", str(e))
            return

        if not detectado:
            tk.Label(self.frm_items,
                     text="No se detectaron enseñanzas con el patrón esperado.",
                     fg="red").grid(row=0, column=0, columnspan=4, pady=10)
            return

        for mat, profes in sorted(detectado.items()):
            self.grupos[mat] = {"materias": {mat}, "profes": set(profes)}
            self.vars_sel[mat] = tk.BooleanVar(value=True)

        self._redibujar()

    # ── Dibujo ────────────────────────────────────────────────────────────────

    def _redibujar(self):
        for w in self.frm_items.winfo_children():
            w.destroy()
        self.filas_ui.clear()

        for i, nombre in enumerate(sorted(self.grupos)):
            g = self.grupos[nombre]
            var = self.vars_sel[nombre]

            tk.Checkbutton(self.frm_items, variable=var).grid(
                row=i, column=0, padx=4, pady=2, sticky="w")
            tk.Label(self.frm_items, text=nombre, anchor="w", width=26,
                     font=("Arial", 9, "bold")).grid(
                row=i, column=1, sticky="w", padx=2)
            tk.Label(self.frm_items,
                     text=", ".join(sorted(g["materias"])),
                     anchor="w", width=28, wraplength=215,
                     justify="left", fg="#555555").grid(
                row=i, column=2, sticky="w", padx=2)
            tk.Label(self.frm_items,
                     text=", ".join(sorted(g["profes"])),
                     anchor="w", width=22, wraplength=180,
                     justify="left").grid(
                row=i, column=3, sticky="w", padx=2)

    # ── Acciones ──────────────────────────────────────────────────────────────

    def sel_todo(self):
        for v in self.vars_sel.values():
            v.set(True)

    def desel_todo(self):
        for v in self.vars_sel.values():
            v.set(False)

    def fusionar(self):
        sel = [n for n, v in self.vars_sel.items() if v.get()]
        if len(sel) < 2:
            messagebox.showwarning(
                "Fusionar", "Marca al menos dos enseñanzas para fusionarlas.")
            return

        # Nombre sugerido: prefijo común de palabras entre los seleccionados
        def prefijo_comun(nombres):
            listas = [n.split() for n in nombres]
            comun = []
            for palabras in zip(*listas):
                if len(set(palabras)) == 1:
                    comun.append(palabras[0])
                else:
                    break
            return " ".join(comun) if comun else nombres[0].split()[0]
        nombre_prop = prefijo_comun(sel)
        nuevo = simpledialog.askstring(
            "Nombre del grupo fusionado",
            "Enseñanzas a fusionar:\n  • " + "\n  • ".join(sorted(sel))
            + "\n\nNombre para el grupo:",
            initialvalue=nombre_prop,
            parent=self.root)
        if not nuevo:
            return
        nuevo = nuevo.strip().upper()

        mats, profes = set(), set()
        for n in sel:
            mats |= self.grupos[n]["materias"]
            profes |= self.grupos[n]["profes"]
            del self.grupos[n]
            del self.vars_sel[n]

        self.grupos[nuevo] = {"materias": mats, "profes": profes}
        self.vars_sel[nuevo] = tk.BooleanVar(value=False)
        self._redibujar()

    def separar(self):
        sel = [n for n, v in self.vars_sel.items() if v.get()]
        if len(sel) != 1:
            messagebox.showwarning(
                "Separar", "Marca exactamente un grupo para separarlo.")
            return
        nombre = sel[0]
        g = self.grupos[nombre]
        if len(g["materias"]) < 2:
            messagebox.showinfo(
                "Separar", "Este grupo tiene una sola asignatura, no se puede separar.")
            return

        mats = sorted(g["materias"])
        del self.grupos[nombre]
        del self.vars_sel[nombre]

        try:
            detectado = extraer_ensenanzas_y_profesores(self.ruta_docx.get())
        except Exception:
            detectado = {}

        for mat in mats:
            self.grupos[mat] = {"materias": {mat},
                                "profes": detectado.get(mat, set())}
            self.vars_sel[mat] = tk.BooleanVar(value=True)

        self._redibujar()

    # ── Generación ────────────────────────────────────────────────────────────

    def copiar_a_grupos(self):
        """Copia las materias del grupo marcado (sin moverlas) a otros grupos."""
        sel = [n for n, v in self.vars_sel.items() if v.get()]
        if not sel:
            messagebox.showwarning("Copiar", "Marca al menos una enseñanza para copiar.")
            return
        destinos_posibles = sorted(g for g in self.grupos if g not in sel)
        if not destinos_posibles:
            messagebox.showwarning("Copiar", "No hay grupos destino disponibles.")
            return

        dlg = tk.Toplevel(self.root)
        dlg.title("Copiar a grupos")
        dlg.grab_set()
        ancho, alto = 400, 360
        px = (self.root.winfo_screenwidth() // 2) - (ancho // 2)
        dlg.geometry(f"{ancho}x{alto}+{px}+100")

        tk.Label(dlg, text="Se copiará el contenido de:",
                 font=("Arial", 9, "bold")).pack(anchor="w", padx=12, pady=(10, 0))
        tk.Label(dlg, text="  • " + "\n  • ".join(sorted(sel)),
                 fg="#1565C0", justify="left").pack(anchor="w", padx=20)

        tk.Label(dlg, text="Selecciona los grupos destino:",
                 font=("Arial", 9, "bold")).pack(anchor="w", padx=12, pady=(10, 0))

        frm_cb = tk.Frame(dlg)
        frm_cb.pack(fill=tk.BOTH, expand=True, padx=20, pady=4)
        vars_dest = {}
        for dest in destinos_posibles:
            v = tk.BooleanVar(value=False)
            vars_dest[dest] = v
            tk.Checkbutton(frm_cb, text=dest, variable=v, anchor="w").pack(fill=tk.X)

        def confirmar():
            destinos = [d for d, v in vars_dest.items() if v.get()]
            if not destinos:
                messagebox.showwarning("Copiar", "Selecciona al menos un destino.",
                                       parent=dlg)
                return
            for origen in sel:
                g_orig = self.grupos[origen]
                for dest in destinos:
                    self.grupos[dest]["materias"] |= g_orig["materias"]
                    self.grupos[dest]["profes"]   |= g_orig["profes"]
            dlg.destroy()
            self._redibujar()

        tk.Button(dlg, text="Copiar", command=confirmar,
                  bg="#1565C0", fg="white",
                  font=("Arial", 10, "bold")).pack(pady=12)

    def generar(self):
        if not self.ruta_docx.get():
            messagebox.showwarning("Aviso", "Primero selecciona el archivo Word.")
            return

        tareas = [
            {"nombre": nombre,
             "materias_id": self.grupos[nombre]["materias"],
             "profes_id": self.grupos[nombre]["profes"],
             "nombre_archivo": f"Horario_{nombre.replace(' ', '_')}"}
            for nombre, var in self.vars_sel.items() if var.get()
        ]
        if not tareas:
            messagebox.showwarning("Aviso", "No hay enseñanzas seleccionadas.")
            return

        out_d = filedialog.askdirectory(title="Carpeta de destino")
        if not out_d:
            return

        try:
            for tarea in tareas:
                doc = Document(self.ruta_docx.get())
                es_espad = "ESPAD" in tarea["nombre"]

                for tabla in doc.tables:
                    for f_idx, fila in enumerate(tabla.rows):
                        # Fila 0 = cabecera con los días de la semana: no tocar
                        if f_idx == 0:
                            continue

                        for c_idx, celda in enumerate(fila.cells):
                            if c_idx == 0:
                                continue

                            for p in celda.paragraphs:
                                texto_p = eliminar_tildes(p.text)
                                if not texto_p:
                                    continue
                                tiene_profe = any(
                                    pr in texto_p for pr in tarea["profes_id"])
                                tiene_mat = any(
                                    mat in texto_p for mat in tarea["materias_id"])

                                if tiene_profe and tiene_mat:
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    if es_espad:
                                        for run in p.runs:
                                            aplicar_color_espad(run, p.text)
                                else:
                                    for run in p.runs:
                                        run.text = ""

                            for p in celda.paragraphs[:]:
                                if not p.text.strip() and len(celda.paragraphs) > 1:
                                    p._element.getparent().remove(p._element)

                            celda.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    # Eliminar filas de datos totalmente vacías (nunca la fila 0)
                    for i in range(len(tabla.rows) - 1, 0, -1):
                        f = tabla.rows[i]
                        if not "".join(
                                c.text.strip()
                                for idx, c in enumerate(f.cells) if idx > 0):
                            f._tr.getparent().remove(f._tr)

                    # Corregir borde inferior de la última fila tras el filtrado
                    from lxml import etree
                    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                    filas_restantes = tabla.rows
                    if filas_restantes:
                        ultima_fila = filas_restantes[-1]
                        for celda in ultima_fila.cells:
                            tcBorders = celda._tc.find(
                                f"{{{WNS}}}tcPr/{{{WNS}}}tcBorders")
                            if tcBorders is not None:
                                bottom = tcBorders.find(f"{{{WNS}}}bottom")
                                if bottom is not None:
                                    # Si el borde inferior estaba anulado, restaurarlo
                                    if bottom.get(f"{{{WNS}}}val") == "nil":
                                        bottom.set(f"{{{WNS}}}val", "single")
                                        bottom.set(f"{{{WNS}}}sz", "4")
                                        bottom.set(f"{{{WNS}}}space", "0")
                                        bottom.set(f"{{{WNS}}}color", "000000")

                doc.save(os.path.join(out_d, f"{tarea['nombre_archivo']}.docx"))

            if self.abrir_carpeta.get():
                import subprocess, sys
                if sys.platform == "win32":
                    os.startfile(out_d)
                elif sys.platform == "darwin":
                    subprocess.Popen(["open", out_d])
                else:
                    subprocess.Popen(["xdg-open", out_d])
            messagebox.showinfo(
                "Finalizado",
                f"Se han generado {len(tareas)} documento(s) en:\n{out_d}")
            self.root.destroy()

        except Exception as e:
            messagebox.showerror("Error", str(e))


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    r = tk.Tk()
    AppAuto(r)
    r.mainloop()
