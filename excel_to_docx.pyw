"""
excel_to_docx.pyw
Convierte una tabla de Excel (.xlsx/.xls/.csv) en una tabla Word (.docx).
GUI centrada, pegada a 5 px del borde superior de pantalla.
Requiere:  pip install openpyxl python-docx pandas
"""

import os
import sys
import subprocess
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading

# ── Importaciones opcionales ─────────────────────────────────────────────────
try:
    import pandas as pd
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    LIBS_OK = True
except ImportError:
    LIBS_OK = False


# ── Colores y estilos ─────────────────────────────────────────────────────────
BG        = "#F4F6F9"
ACCENT    = "#2563EB"
ACCENT_H  = "#1D4ED8"
BTN_FG    = "#FFFFFF"
ENTRY_BG  = "#FFFFFF"
BORDER    = "#CBD5E1"
TEXT_MAIN = "#1E293B"
TEXT_SUB  = "#64748B"
SUCCESS   = "#16A34A"
ERROR_CLR = "#DC2626"


def set_cell_shading(cell, fill_hex: str):
    """Aplica color de fondo a una celda de tabla Word."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.upper())
    tcPr.append(shd)


def excel_to_docx(xlsx_path: str,
                  docx_path: str,
                  sheet_idx: int,
                  header_color: str,
                  header_text_color: str,
                  font_size: int,
                  alternate_rows: bool) -> None:
    """Lee el Excel y escribe el DOCX con la tabla formateada."""
    df = pd.read_excel(xlsx_path, sheet_name=sheet_idx, header=0, dtype=str)
    df.fillna("", inplace=True)

    doc   = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(font_size)

    # Márgenes de página (2 cm)
    for sec in doc.sections:
        sec.left_margin   = Inches(0.8)
        sec.right_margin  = Inches(0.8)
        sec.top_margin    = Inches(0.8)
        sec.bottom_margin = Inches(0.8)

    cols = list(df.columns)
    ncols = len(cols)
    nrows = len(df)

    table = doc.add_table(rows=1 + nrows, cols=ncols)
    table.style = "Table Grid"

    # ── Encabezado ────────────────────────────────────────────────────────────
    hdr_row = table.rows[0]
    for i, col in enumerate(cols):
        cell = hdr_row.cells[i]
        cell.text = str(col)
        set_cell_shading(cell, header_color.lstrip("#"))
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = para.runs[0]
        run.bold       = True
        run.font.size  = Pt(font_size)
        # color del texto de cabecera
        r, g, b = tuple(int(header_text_color.lstrip("#")[j:j+2], 16) for j in (0, 2, 4))
        run.font.color.rgb = RGBColor(r, g, b)

    # ── Filas de datos ────────────────────────────────────────────────────────
    ALT_COLOR = "EFF6FF"   # azul muy claro para filas alternas
    for row_i, (_, serie) in enumerate(df.iterrows()):
        tr = table.rows[row_i + 1]
        bg = ALT_COLOR if (alternate_rows and row_i % 2 == 1) else "FFFFFF"
        for col_i, val in enumerate(serie):
            cell = tr.cells[col_i]
            cell.text = str(val)
            set_cell_shading(cell, bg)
            para = cell.paragraphs[0]
            run  = para.runs[0] if para.runs else para.add_run(str(val))
            run.font.size = Pt(font_size)

    # Ancho uniforme de columnas
    total_w = Inches(8.7 / max(ncols, 1))
    for row in table.rows:
        for cell in row.cells:
            cell.width = total_w

    doc.save(docx_path)


# ══════════════════════════════════════════════════════════════════════════════
#  GUI
# ══════════════════════════════════════════════════════════════════════════════
class App(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title("Excel → DOCX  |  Tabla Word")
        self.resizable(False, False)
        self.configure(bg=BG)

        # ── Variables ─────────────────────────────────────────────────────────
        self.xlsx_var         = tk.StringVar()
        self.docx_var         = tk.StringVar()
        self.sheet_var        = tk.StringVar(value="0")
        self.hdr_color_var    = tk.StringVar(value="#2563EB")
        self.hdr_txt_var      = tk.StringVar(value="#FFFFFF")
        self.font_size_var    = tk.IntVar(value=10)
        self.alt_rows_var     = tk.BooleanVar(value=True)
        self.open_folder_var  = tk.BooleanVar(value=True)

        self._build_ui()
        self._center_top()   # centrar horizontalmente, 5 px del borde superior

    # ── Layout ───────────────────────────────────────────────────────────────
    def _build_ui(self):
        pad = dict(padx=18, pady=6)

        # Título
        tk.Label(self, text="Excel  →  Tabla Word (.docx)",
                 bg=BG, fg=ACCENT,
                 font=("Segoe UI", 15, "bold")).grid(row=0, column=0, columnspan=3,
                                                     padx=18, pady=(18, 4), sticky="w")
        tk.Label(self, text="Convierte cualquier hoja de Excel en una tabla formateada en Word",
                 bg=BG, fg=TEXT_SUB,
                 font=("Segoe UI", 9)).grid(row=1, column=0, columnspan=3,
                                            padx=18, pady=(0, 14), sticky="w")

        # ── Archivo origen ────────────────────────────────────────────────────
        self._label(2, "Archivo Excel (.xlsx / .xls / .csv)")
        self._entry_btn(3, self.xlsx_var, self._browse_xlsx, "Examinar…")

        # Hoja
        self._label(4, "Número de hoja  (0 = primera, 1 = segunda, …)")
        tk.Entry(self, textvariable=self.sheet_var, width=6,
                 bg=ENTRY_BG, relief="flat",
                 highlightbackground=BORDER,
                 highlightthickness=1,
                 font=("Segoe UI", 10)).grid(row=5, column=0, padx=18, pady=4, sticky="w")

        # ── Archivo destino ───────────────────────────────────────────────────
        self._label(6, "Guardar DOCX en…")
        self._entry_btn(7, self.docx_var, self._browse_docx, "Guardar como…")

        # ── Opciones de formato ───────────────────────────────────────────────
        sep = ttk.Separator(self, orient="horizontal")
        sep.grid(row=8, column=0, columnspan=3, sticky="ew", padx=18, pady=10)

        tk.Label(self, text="Opciones de formato",
                 bg=BG, fg=TEXT_MAIN,
                 font=("Segoe UI", 10, "bold")).grid(row=9, column=0, columnspan=3,
                                                     padx=18, pady=(0, 6), sticky="w")

        # Color cabecera + color texto cabecera + tamaño fuente — fila 10
        frm_opts = tk.Frame(self, bg=BG)
        frm_opts.grid(row=10, column=0, columnspan=3, padx=18, pady=4, sticky="w")

        tk.Label(frm_opts, text="Color cabecera:", bg=BG, fg=TEXT_MAIN,
                 font=("Segoe UI", 9)).grid(row=0, column=0, sticky="w", padx=(0, 4))
        self.hdr_preview = tk.Label(frm_opts, width=3, bg=self.hdr_color_var.get(),
                                    relief="flat", bd=1)
        self.hdr_preview.grid(row=0, column=1, padx=(0, 4))
        tk.Entry(frm_opts, textvariable=self.hdr_color_var, width=9,
                 bg=ENTRY_BG, relief="flat",
                 highlightbackground=BORDER, highlightthickness=1,
                 font=("Segoe UI", 9)).grid(row=0, column=2, padx=(0, 16))
        self.hdr_color_var.trace_add("write", self._update_hdr_preview)

        tk.Label(frm_opts, text="Color texto:", bg=BG, fg=TEXT_MAIN,
                 font=("Segoe UI", 9)).grid(row=0, column=3, sticky="w", padx=(0, 4))
        self.txt_preview = tk.Label(frm_opts, width=3, bg=self.hdr_txt_var.get(),
                                    relief="flat", bd=1)
        self.txt_preview.grid(row=0, column=4, padx=(0, 4))
        tk.Entry(frm_opts, textvariable=self.hdr_txt_var, width=9,
                 bg=ENTRY_BG, relief="flat",
                 highlightbackground=BORDER, highlightthickness=1,
                 font=("Segoe UI", 9)).grid(row=0, column=5, padx=(0, 16))
        self.hdr_txt_var.trace_add("write", self._update_txt_preview)

        tk.Label(frm_opts, text="Tamaño fuente:", bg=BG, fg=TEXT_MAIN,
                 font=("Segoe UI", 9)).grid(row=0, column=6, sticky="w", padx=(0, 4))
        tk.Spinbox(frm_opts, from_=7, to=20, textvariable=self.font_size_var,
                   width=4, font=("Segoe UI", 9)).grid(row=0, column=7)

        # Alternar filas
        tk.Checkbutton(self, text="Alternar color en filas pares",
                       variable=self.alt_rows_var,
                       bg=BG, fg=TEXT_MAIN,
                       activebackground=BG,
                       font=("Segoe UI", 9),
                       selectcolor=ENTRY_BG).grid(row=11, column=0, columnspan=3,
                                                   padx=18, pady=4, sticky="w")

        # Abrir carpeta al terminar
        tk.Checkbutton(self, text="Abrir carpeta de destino al finalizar",
                       variable=self.open_folder_var,
                       bg=BG, fg=TEXT_MAIN,
                       activebackground=BG,
                       font=("Segoe UI", 9),
                       selectcolor=ENTRY_BG).grid(row=12, column=0, columnspan=3,
                                                   padx=18, pady=4, sticky="w")

        # ── Barra de progreso + estado ────────────────────────────────────────
        self.progress = ttk.Progressbar(self, mode="indeterminate", length=380)
        self.progress.grid(row=13, column=0, columnspan=3, padx=18, pady=(12, 2), sticky="ew")

        self.status_var = tk.StringVar(value="Listo.")
        tk.Label(self, textvariable=self.status_var,
                 bg=BG, fg=TEXT_SUB,
                 font=("Segoe UI", 9),
                 anchor="w").grid(row=14, column=0, columnspan=3,
                                   padx=18, pady=(0, 4), sticky="ew")

        # ── Botón Convertir ───────────────────────────────────────────────────
        self.btn_convert = tk.Button(self,
                                     text="⚡  Convertir a DOCX",
                                     command=self._start,
                                     bg=ACCENT, fg=BTN_FG,
                                     activebackground=ACCENT_H,
                                     activeforeground=BTN_FG,
                                     relief="flat", bd=0,
                                     font=("Segoe UI", 11, "bold"),
                                     padx=22, pady=10,
                                     cursor="hand2")
        self.btn_convert.grid(row=15, column=0, columnspan=3,
                               padx=18, pady=(6, 18), sticky="ew")

        self.columnconfigure(0, weight=1)

    # ── Helpers UI ────────────────────────────────────────────────────────────
    def _label(self, row, text):
        tk.Label(self, text=text,
                 bg=BG, fg=TEXT_MAIN,
                 font=("Segoe UI", 9)).grid(row=row, column=0, columnspan=3,
                                             padx=18, pady=(8, 0), sticky="w")

    def _entry_btn(self, row, var, cmd, btn_text):
        frame = tk.Frame(self, bg=BG)
        frame.grid(row=row, column=0, columnspan=3, padx=18, pady=4, sticky="ew")
        e = tk.Entry(frame, textvariable=var, width=46,
                     bg=ENTRY_BG, relief="flat",
                     highlightbackground=BORDER, highlightthickness=1,
                     font=("Segoe UI", 9))
        e.pack(side="left", padx=(0, 8))
        tk.Button(frame, text=btn_text, command=cmd,
                  bg=BORDER, fg=TEXT_MAIN,
                  relief="flat", bd=0,
                  font=("Segoe UI", 9),
                  padx=10, pady=4,
                  cursor="hand2").pack(side="left")

    def _update_hdr_preview(self, *_):
        try:
            self.hdr_preview.configure(bg=self.hdr_color_var.get())
        except Exception:
            pass

    def _update_txt_preview(self, *_):
        try:
            self.txt_preview.configure(bg=self.hdr_txt_var.get())
        except Exception:
            pass

    # ── Posición: centrada, 5 px del borde superior ───────────────────────────
    def _center_top(self):
        self.update_idletasks()
        w = self.winfo_reqwidth()
        sw = self.winfo_screenwidth()
        x  = (sw - w) // 2
        self.geometry(f"+{x}+5")

    # ── Diálogos de archivo ───────────────────────────────────────────────────
    def _browse_xlsx(self):
        path = filedialog.askopenfilename(
            title="Seleccionar archivo Excel",
            filetypes=[("Excel / CSV", "*.xlsx *.xls *.csv"), ("Todos", "*.*")])
        if path:
            self.xlsx_var.set(path)
            # Proponer nombre de salida automático
            base = os.path.splitext(path)[0]
            self.docx_var.set(base + "_tabla.docx")

    def _browse_docx(self):
        path = filedialog.asksaveasfilename(
            title="Guardar DOCX como…",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")])
        if path:
            self.docx_var.set(path)

    # ── Conversión ────────────────────────────────────────────────────────────
    def _start(self):
        xlsx = self.xlsx_var.get().strip()
        docx = self.docx_var.get().strip()

        if not xlsx:
            messagebox.showwarning("Falta archivo", "Selecciona el archivo Excel de origen.")
            return
        if not os.path.isfile(xlsx):
            messagebox.showerror("No encontrado", f"No se encuentra:\n{xlsx}")
            return
        if not docx:
            messagebox.showwarning("Falta destino", "Indica dónde guardar el archivo DOCX.")
            return

        try:
            sheet_idx = int(self.sheet_var.get())
        except ValueError:
            messagebox.showerror("Error", "El número de hoja debe ser un entero (0, 1, 2…).")
            return

        self.btn_convert.configure(state="disabled")
        self.progress.start(10)
        self.status_var.set("Convirtiendo…  ⏳")
        self.update_idletasks()

        threading.Thread(target=self._run,
                         args=(xlsx, docx, sheet_idx),
                         daemon=True).start()

    def _run(self, xlsx, docx, sheet_idx):
        try:
            excel_to_docx(
                xlsx_path        = xlsx,
                docx_path        = docx,
                sheet_idx        = sheet_idx,
                header_color     = self.hdr_color_var.get(),
                header_text_color= self.hdr_txt_var.get(),
                font_size        = self.font_size_var.get(),
                alternate_rows   = self.alt_rows_var.get(),
            )
            self.after(0, self._done, docx)
        except Exception as exc:
            self.after(0, self._error, str(exc))

    def _done(self, docx_path):
        self.progress.stop()
        self.btn_convert.configure(state="normal")
        self.status_var.set(f"✅  ¡Listo!  →  {os.path.basename(docx_path)}")
        messagebox.showinfo("Conversión completada",
                            f"Archivo guardado:\n{docx_path}")
        if self.open_folder_var.get():
            folder = os.path.dirname(os.path.abspath(docx_path))
            if sys.platform == "win32":
                os.startfile(folder)
            elif sys.platform == "darwin":
                subprocess.Popen(["open", folder])
            else:
                subprocess.Popen(["xdg-open", folder])

    def _error(self, msg):
        self.progress.stop()
        self.btn_convert.configure(state="normal")
        self.status_var.set("❌  Error durante la conversión.")
        messagebox.showerror("Error", f"No se pudo convertir:\n\n{msg}")


# ══════════════════════════════════════════════════════════════════════════════
#  Comprobación de dependencias al arrancar
# ══════════════════════════════════════════════════════════════════════════════
def check_deps():
    missing = []
    for pkg in ("pandas", "openpyxl", "docx"):
        try:
            __import__(pkg)
        except ImportError:
            missing.append(pkg)
    if missing:
        root = tk.Tk()
        root.withdraw()
        install = messagebox.askyesno(
            "Dependencias faltantes",
            f"Faltan los paquetes:\n  {', '.join(missing)}\n\n"
            "¿Instalarlos ahora con pip?")
        root.destroy()
        if install:
            subprocess.check_call([sys.executable, "-m", "pip", "install",
                                   "pandas", "openpyxl", "python-docx"])
        else:
            sys.exit(1)


if __name__ == "__main__":
    check_deps()
    app = App()
    app.mainloop()
