import fitz  # PyMuPDF
import pandas as pd
from docx import Document
import re, os, subprocess
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

class AppExamenesPro:
    def __init__(self, root):
        self.root = root
        self.root.title("Generador de Exámenes")
        ancho = 700
        alto = 700
        pos_x = (self.root.winfo_screenwidth() // 2) - (ancho // 2)
        self.root.geometry(f"{ancho}x{alto}+{pos_x}+15")
        
        # Variable para guardar el nombre del excel cargado
        self.nombre_archivo_excel = "Pro" 

        self.tab_control = ttk.Notebook(root)
        self.tab1 = ttk.Frame(self.tab_control)
        self.tab2 = ttk.Frame(self.tab_control)
        self.tab_control.add(self.tab1, text=' 1. Extraer preguntas de PDF a Excel ')
        self.tab_control.add(self.tab2, text=' 2. Crear Examen en Word ')
        self.tab_control.pack(expand=1, fill="both")

        self.setup_tab1()
        self.setup_tab2()

    def setup_tab1(self):
        self.ruta_pdf = tk.StringVar()
        self.ruta_excel_dest = tk.StringVar()
        self.nombre_excel_sugerido = ""
        
        tk.Label(self.tab1, text="PASO A: Convertir apuntes en base de datos de preguntas", font=("Arial", 12, "bold")).pack(pady=15)
        
        frame_pdf = tk.LabelFrame(self.tab1, text=" 1. Selecciona el PDF original ", padx=10, pady=10)
        frame_pdf.pack(fill="x", padx=20)
        
        tk.Entry(frame_pdf, textvariable=self.ruta_pdf, width=65).pack(side=tk.LEFT, padx=5)
        tk.Button(frame_pdf, text="Buscar...", command=self.seleccionar_pdf).pack(side=tk.LEFT)

        # --- NUEVA ETIQUETA PARA MOSTRAR PÁGINAS ---
        self.lbl_info_pdf = tk.Label(self.tab1, text="No se ha cargado ningún PDF", font=("Arial", 9, "italic"), fg="gray")
        self.lbl_info_pdf.pack(pady=2)
        # ------------------------------------------

        tk.Label(self.tab1, text="2. Páginas a extraer (ej: 7, 9, 15-20):", font=("Arial", 10)).pack(pady=10)
        self.ent_pags = tk.Entry(self.tab1, width=40)
        self.ent_pags.pack()
        
        # ... resto del código del setup_tab1 ...
        frame_ex = tk.LabelFrame(self.tab1, text=" 3. Guardar Almacén (Excel) como... ", padx=10, pady=10)
        frame_ex.pack(fill="x", padx=20, pady=15)
        tk.Entry(frame_ex, textvariable=self.ruta_excel_dest, width=65).pack(side=tk.LEFT, padx=5)
        tk.Button(frame_ex, text="Carpeta...", command=self.seleccionar_destino_excel).pack(side=tk.LEFT)
        tk.Button(self.tab1, text="EJECUTAR EXTRACCIÓN", command=self.run_extraccion, bg="#4CAF50", fg="white", font=("Arial", 11, "bold"), height=2).pack(pady=20)

    def seleccionar_pdf(self):
        f = filedialog.askopenfilename(filetypes=[("PDF", "*.pdf")])
        if f:
            self.ruta_pdf.set(f)
            nombre = os.path.splitext(os.path.basename(f))[0]
            self.nombre_excel_sugerido = f"Banco_Preguntas_{nombre}.xlsx"
            self.ruta_excel_dest.set(self.nombre_excel_sugerido)
            
            try:
                doc = fitz.open(f)
                self.total_paginas_pdf = len(doc) # Guardamos el total para validar después
                self.lbl_info_pdf.config(text=f"El PDF tiene {self.total_paginas_pdf} páginas.", fg="blue")
                doc.close()
            except Exception as e:
                self.lbl_info_pdf.config(text="Error al leer el PDF", fg="red")

    def run_extraccion(self):
        try:
            ruta = self.ruta_pdf.get()
            if not ruta:
                messagebox.showwarning("Atención", "Selecciona un PDF primero.")
                return

            doc_pdf = fitz.open(ruta)
            pags_str = self.ent_pags.get()
            indices = []
            
            # --- MEJORA 1: VALIDACIÓN DE PÁGINAS ---
            try:
                for p in pags_str.split(','):
                    if '-' in p:
                        ini, fin = map(int, p.split('-'))
                        if ini < 1 or fin > self.total_paginas_pdf:
                            raise ValueError(f"Rango {p} fuera de límites (1-{self.total_paginas_pdf})")
                        indices.extend(range(ini-1, fin))
                    else:
                        val = int(p.strip())
                        if val < 1 or val > self.total_paginas_pdf:
                            raise ValueError(f"Página {val} no existe")
                        indices.append(val-1)
            except ValueError as ve:
                messagebox.showerror("Error de Rango", f"Página no válida: {ve}")
                return

            preguntas = []
            pregunta_actual = ""
            
            # --- MEJORA 2: REGEX FILTRADO (Excluye 1.1, 2.1.3, etc.) ---
            # Explicación: ^\d+ significa que empiece por números.
            # [\.\)] significa que le siga un punto o paréntesis.
            # (?!\d) es un "negative lookahead": asegura que después del punto NO haya otro número.
            patron_pregunta = r'^\d+[\.\)](?!\d)'

            for idx in indices:
                if idx >= len(doc_pdf): continue
                for b in doc_pdf[idx].get_text("dict")["blocks"]:
                    if "lines" not in b: continue
                    for l in b["lines"]:
                        txt = "".join([s["text"] for s in l["spans"]]).strip()
                        bold = any(s["flags"] & 2 or "bold" in s["font"].lower() for s in l["spans"])
                        
                        # Aplicamos el nuevo filtro
                        if bold and re.match(patron_pregunta, txt):
                            if pregunta_actual: preguntas.append(pregunta_actual.strip())
                            pregunta_actual = txt
                        elif pregunta_actual and "entregar" not in txt.lower():
                            pregunta_actual += " " + txt
            
            if pregunta_actual: preguntas.append(pregunta_actual.strip())
            
            if not preguntas:
                messagebox.showwarning("Aviso", "No se detectaron preguntas con negrita y numeración simple.")
                return

            pd.DataFrame({"Pregunta": preguntas}).to_excel(self.ruta_excel_dest.get(), index=False)
            messagebox.showinfo("Éxito", f"Excel creado con {len(preguntas)} preguntas.")
            
        except Exception as e: 
            messagebox.showerror("Error", f"Fallo al extraer: {e}")

    def seleccionar_destino_excel(self):
        f = filedialog.asksaveasfilename(initialfile=self.nombre_excel_sugerido, defaultextension=".xlsx")
        if f: self.ruta_excel_dest.set(f)

    def setup_tab2(self):
        self.preguntas_db = []
        self.vars_checks = []
        self.lbl_contador = tk.StringVar(value="Seleccionadas: 0 de 6")
        
        tk.Label(self.tab2, text="PASO B: Seleccionar preguntas y crear examen", font=("Arial", 12, "bold")).pack(pady=10)
        
        frame_config = tk.Frame(self.tab2)
        frame_config.pack(pady=5)
        
        btn_load = tk.Button(frame_config, text="Cargar Almacén (Excel)", command=self.load_excel, bg="#f0f0f0", padx=10)
        btn_load.pack(side=tk.LEFT, padx=10)
        
        tk.Label(frame_config, text="Nº de preguntas en plantilla:").pack(side=tk.LEFT)
        self.ent_max_preguntas = tk.Entry(frame_config, width=5)
        self.ent_max_preguntas.insert(0, "6")
        self.ent_max_preguntas.pack(side=tk.LEFT, padx=5)
        self.ent_max_preguntas.bind("<KeyRelease>", self.actualizar_conteo)

        self.display_conteo = tk.Label(self.tab2, textvariable=self.lbl_contador, font=("Arial", 11, "bold"), fg="blue")
        self.display_conteo.pack(pady=5)

        self.canvas_frame = tk.Frame(self.tab2, bd=2, relief="sunken")
        self.canvas_frame.pack(fill="both", expand=True, padx=20)
        
        self.canvas = tk.Canvas(self.canvas_frame)
        sb = ttk.Scrollbar(self.canvas_frame, orient="vertical", command=self.canvas.yview)
        self.scroll_list = tk.Frame(self.canvas)
        self.scroll_list.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.create_window((0, 0), window=self.scroll_list, anchor="nw")
        self.canvas.configure(yscrollcommand=sb.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")

        self.btn_gen_word = tk.Button(self.tab2, text="GENERAR EXAMEN(WORD)", command=self.create_word, bg="#2196F3", fg="white", font=("Arial", 11, "bold"), state="disabled", height=2)
        self.btn_gen_word.pack(pady=20)

    def load_excel(self):
        f = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx")])
        if f:
            # --- MEJORA: Guardamos el nombre del excel para el futuro ---
            self.nombre_archivo_excel = os.path.splitext(os.path.basename(f))[0]
            
            df = pd.read_excel(f)
            self.preguntas_db = df['Pregunta'].tolist()
            self.actualizar_lista()
            self.btn_gen_word.config(state="normal")

    def actualizar_lista(self):
        for w in self.scroll_list.winfo_children(): w.destroy()
        self.vars_checks = []
        for i, p in enumerate(self.preguntas_db):
            v = tk.BooleanVar()
            v.trace_add("write", self.actualizar_conteo)
            self.vars_checks.append(v)
            texto_chk = f"{i+1}. {p[:120]}..." if len(p) > 120 else f"{i+1}. {p}"
            tk.Checkbutton(self.scroll_list, text=texto_chk, variable=v, justify=tk.LEFT, wraplength=600, anchor="w").pack(fill="x", padx=5, pady=2)

    def actualizar_conteo(self, *args):
        try:
            limite = int(self.ent_max_preguntas.get())
        except:
            limite = 0
        seleccionadas = sum(1 for v in self.vars_checks if v.get())
        self.lbl_contador.set(f"Seleccionadas: {seleccionadas} de {limite}")
        self.display_conteo.config(fg="red" if seleccionadas > limite else "blue")

    def maquetar_texto(self, texto):
        texto = re.sub(r'\?', '?\n', texto)
        texto = re.sub(r'\s+([a-zA-Z]\))', r'\n\1', texto)
        texto = re.sub(r'\s+•', r'\n•', texto)
        texto = re.sub(r'\.\s+([A-Z]\.)', r'.\n\1', texto)
        texto = re.sub(r'\.\s+(_+)', r'.\n\1', texto)
        texto = re.sub(r'\.\s+(\d+)\s+', r'.\n\1 ', texto)
        texto = re.sub(r'\s+(\d+[\.\)])', r'\n\1', texto)
        texto = re.sub(r'\n+', '\n', texto)
        return texto.strip()

    def create_word(self):
        try:
            limite = int(self.ent_max_preguntas.get())
        except:
            messagebox.showerror("Error", "Introduce un número válido.")
            return

        sel = [self.preguntas_db[i] for i, v in enumerate(self.vars_checks) if v.get()]
        if not sel: return
        
        plantilla = filedialog.askopenfilename(title="Elige Plantilla examen.docx", filetypes=[("Word", "*.docx")])
        if not plantilla: return
        
        # --- MEJORA: Aquí toma el nombre del Excel ---
        sugerencia_nombre = f"Examen_{self.nombre_archivo_excel}.docx"
        out = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=sugerencia_nombre)
        
        if not out: return

        try:
            doc = Document(plantilla)
            n = 1
            for p in doc.paragraphs:
                marcador = f"Pregunta {n}."
                if marcador in p.text and n <= len(sel) and n <= limite:
                    texto_limpio = self.maquetar_texto(sel[n-1])
                    p.add_run(f" {texto_limpio}")
                    n += 1
            doc.save(out)
            messagebox.showinfo("¡Éxito!", "Examen guardado.")
            os.startfile(os.path.dirname(out))
        except Exception as e: messagebox.showerror("Error", f"Fallo: {e}")

if __name__ == "__main__":
    root = tk.Tk()
    AppExamenesPro(root)
    root.mainloop()