import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import unicodedata

def eliminar_tildes(cadena):
    if not cadena: return ""
    s = "".join(c for c in unicodedata.normalize('NFD', str(cadena))
               if unicodedata.category(c) != 'Mn')
    return s.upper().strip()

def aplicar_color_espad(run, texto):
    texto_limpio = eliminar_tildes(texto)
    if "1.2" in texto_limpio:
        run.font.color.rgb = RGBColor(255, 0, 0)
    elif "2.1" in texto_limpio:
        run.font.color.rgb = RGBColor(0, 128, 0)
    elif "2.2" in texto_limpio:
        run.font.color.rgb = RGBColor(0, 0, 255)
    else:
        run.font.color.rgb = RGBColor(128, 0, 128)

class AppFinal:
    def __init__(self, root):
        self.root = root
        self.root.title("Filtro")
        self.root.geometry("600x550")
        self.tareas = []

        tk.Label(root, text="Generador de Horarios por Enseñanzas", font=('Arial', 12, 'bold')).pack(pady=10)
        
        input_f = tk.LabelFrame(root, text=" Configuración ", padx=10, pady=10)
        input_f.pack(fill=tk.X, padx=20, pady=10)

        tk.Label(input_f, text="Materia Clave:").pack(anchor="w")
        self.ent_m = tk.Entry(input_f, width=50)
        self.ent_m.pack(pady=2)

        tk.Label(input_f, text="Profesores (separados por comas):").pack(anchor="w", pady=(10,0))
        self.ent_p = tk.Entry(input_f, width=50)
        self.ent_p.pack(pady=2)

        tk.Button(input_f, text="+ Añadir a la lista", command=self.add, bg="#0078D7", fg="white").pack(pady=15)

        self.lb = tk.Listbox(root, width=75, height=8)
        self.lb.pack(pady=5, padx=20)

        tk.Button(root, text="GENERAR HORARIOS", command=self.run, bg="#2E7D32", fg="white", font=('Arial', 10, 'bold')).pack(pady=20)

    def add(self):
        m, p = self.ent_m.get().strip(), self.ent_p.get().strip()
        if m and p:
            tarea = {
                'materia_id': eliminar_tildes(m),
                'profes_id': [eliminar_tildes(pr.strip()) for pr in p.split(",")],
                'nombre_archivo': f"Horario_{m.replace(' ', '_')}"
            }
            self.tareas.append(tarea)
            self.lb.insert(tk.END, f"📘 {tarea['nombre_archivo']} | Filtro: {m} + {p}")
            self.ent_m.delete(0, tk.END); self.ent_p.delete(0, tk.END)

    def run(self):
        if not self.tareas: return
        file_m = filedialog.askopenfilename(title="Selecciona Word Maestro", filetypes=[("Word", "*.docx")])
        if not file_m: return
        out_d = filedialog.askdirectory(title="Carpeta de destino")
        if not out_d: return

        try:
            for tarea in self.tareas:
                doc = Document(file_m)
                es_espad = "ESPAD" in tarea['materia_id']

                for tabla in doc.tables:
                    for fila in tabla.rows:
                        for c_idx, celda in enumerate(fila.cells):
                            if c_idx == 0: continue
                            
                            # Procesar párrafos
                            pars = celda.paragraphs
                            for p in pars:
                                texto_p = eliminar_tildes(p.text)
                                if texto_p:
                                    tiene_profe = any(profe in texto_p for profe in tarea['profes_id'])
                                    tiene_mat = (tarea['materia_id'] in texto_p)
                                    
                                    if tiene_profe and tiene_mat:
                                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        if es_espad:
                                            for run in p.runs:
                                                aplicar_color_espad(run, p.text)
                                    else:
                                        # Vaciar texto de párrafos que no cumplen el filtro
                                        for run in p.runs:
                                            run.text = ""
                            
                            # Limpieza segura de párrafos vacíos para evitar errores de Word
                            # Dejamos siempre al menos uno (aunque esté vacío) si la celda queda sin nada
                            for p in celda.paragraphs[:]:
                                if not p.text.strip() and len(celda.paragraphs) > 1:
                                    p_element = p._element
                                    p_element.getparent().remove(p_element)
                            
                            celda.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    # Eliminar filas totalmente vacías
                    for i in range(len(tabla.rows) - 1, 0, -1):
                        f_check = tabla.rows[i]
                        contenido = "".join([c.text.strip() for idx, c in enumerate(f_check.cells) if idx > 0])
                        if not contenido:
                            tr = f_check._tr
                            tr.getparent().remove(tr)

                doc.save(os.path.join(out_d, f"{tarea['nombre_archivo']}.docx"))
            
            messagebox.showinfo("Finalizado", "Ya se han generado los documentos.")
            self.root.destroy()
        except Exception as e:
            messagebox.showerror("Error", str(e))

if __name__ == "__main__":
    r = tk.Tk(); app = AppFinal(r); r.mainloop()