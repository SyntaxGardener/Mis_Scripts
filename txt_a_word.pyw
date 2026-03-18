"""
Convertidor TXT → DOCX
=======================
Convierte los archivos .txt generados por el Resumidor al formato Word (.docx).

Para exámenes puede usar una plantilla .docx existente e inyectar el contenido
en ella, actualizando automáticamente los campos en rojo (Módulo y nivel ESPAD)
a partir del nombre del archivo TXT.

Requisito:
    pip install python-docx
"""

import tkinter as tk
from tkinter import ttk, filedialog
import os, sys, re, subprocess
from pathlib import Path
from datetime import datetime
from copy import deepcopy

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Tema visual ───────────────────────────────────────────────────────────────
BG       = "#f4f4f8"
PANEL    = "#e8e8f0"
ENTRADA  = "#ffffff"
ACENTO   = "#5b4fcf"
ACENTO2  = "#2976d4"
TEXTO    = "#1a1a2e"
SUBTEXTO = "#666688"
BORDE    = "#ccccdd"
F_NORM   = ("Segoe UI", 10)
F_TITULO = ("Segoe UI", 13, "bold")
F_LABEL  = ("Segoe UI", 9)
F_BOLD   = ("Segoe UI", 9, "bold")
F_MONO   = ("Consolas", 10)

COLOR_ROJO = RGBColor(0xEE, 0x00, 0x00)   # rojo de la plantilla


# ── Extracción de metadatos desde el nombre del archivo ───────────────────────

def extraer_meta_de_nombre(nombre: str) -> dict:
    """
    Intenta extraer módulo y nivel ESPAD del nombre del archivo TXT.
    Patrones reconocidos (no sensible a mayúsculas):
      - "espad 1.2", "espad1.2", "espad_1.2"
      - "modulo lengua", "mod_lengua_diversidad"
    Devuelve dict con claves 'modulo' y 'espad' (cadenas, pueden estar vacías).
    """
    nombre_limpio = nombre.replace('_', ' ').replace('-', ' ')

    # Nivel ESPAD: busca patron "espad X.Y" o "espad XY"
    m_espad = re.search(r'espad\s*(\d[\d\.]*)', nombre_limpio, re.IGNORECASE)
    espad = m_espad.group(1) if m_espad else ""

    # Módulo: toma el texto antes del primer "_examen", "_resumen", fecha, o espad
    modulo = nombre_limpio
    for patron in [r'\bespad\b.*', r'\d{8}_\d{6}', r'_examen.*', r'_resumen.*',
                   r'_esquema.*', r'_puntos.*', r'_cuestionario.*']:
        modulo = re.split(patron, modulo, flags=re.IGNORECASE)[0]
    modulo = modulo.strip().title()

    return {'modulo': modulo, 'espad': espad}


# ── Detección del tipo de documento ──────────────────────────────────────────

def detectar_tipo(lineas: list) -> str:
    texto = "\n".join(lineas[:10]).upper()
    if "EXAMEN"       in texto: return "examen"
    if "CUESTIONARIO" in texto: return "cuestionario"
    if "ESQUEMA"      in texto: return "esquema"
    if "PUNTOS CLAVE" in texto: return "puntos_clave"
    return "resumen"


# ── Helpers de formato ────────────────────────────────────────────────────────

def es_cabecera_seccion(linea: str) -> bool:
    l = linea.strip()
    if l.startswith(("•", "►", "-", "·", "*")): return False
    if re.match(r'^\d+[).]\s+[A-ZÁÉÍÓÚÜÑ\s/]+$', l):              return True
    if re.match(r'^[A-ZÁÉÍÓÚÜÑ\s/]{6,}:?\s*$', l) and len(l)<60:  return True
    return False

def es_clave_respuestas(linea: str) -> bool:
    kw = ["CLAVE DE RESPUESTAS", "RESPUESTAS CORRECTAS", "SOLUCIONES", "ANSWER KEY"]
    return any(k in linea.upper() for k in kw)

def es_separador(linea: str) -> bool:
    return linea.strip() and all(c in "═─=─-" for c in linea.strip())

def es_opcion_multiple(linea: str) -> bool:
    return bool(re.match(r'^\s*[A-Da-d][).]\s+', linea))

def es_pregunta_numerada(linea: str) -> bool:
    return bool(re.match(r'^\s*\d+[).]\s+', linea))

def añadir_pie(section, texto_izq: str, texto_der: str):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for texto, tab_antes in [(texto_izq, False), (texto_der, True)]:
        if tab_antes:
            rt = p.add_run("\t")
            rt.font.size = Pt(8)
        if texto:
            r = p.add_run(texto)
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x88, 0x8a, 0xaa)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_d = OxmlElement("w:tab")
    tab_d.set(qn("w:val"), "right")
    tab_d.set(qn("w:pos"), "9360")
    tabs.append(tab_d)
    pPr.append(tabs)


# ── Conversión SIN plantilla ──────────────────────────────────────────────────

def convertir_sin_plantilla(ruta_txt: Path, ruta_docx: Path, autor: str = ""):
    doc = Document()

    # Fuente Arial global
    for style in doc.styles:
        try: style.font.name = 'Arial'
        except Exception: pass
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10.5)

    # Márgenes
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3)
        sec.right_margin  = Cm(3)

    lineas = ruta_txt.read_text(encoding='utf-8').splitlines()
    tipo = detectar_tipo(lineas)

    # Título
    titulo_doc = ruta_txt.stem.replace("_", " ").title()
    for linea in lineas:
        l = linea.strip()
        if l and not es_separador(l) and not l.startswith(("✔", "✘")):
            for pref in ["RESUMEN:", "ESQUEMA:", "EXAMEN:", "CUESTIONARIO:", "PUNTOS CLAVE:"]:
                if l.upper().startswith(pref):
                    l = l[len(pref):].strip()
            if l:
                titulo_doc = l
                break

    if tipo == "examen":
        p_t = doc.add_heading("Examen de Comunicación", level=0)
    else:
        p_t = doc.add_heading(titulo_doc, level=0)
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_t.runs:
        p_t.runs[0].font.color.rgb = RGBColor(0x5b, 0x4f, 0xcf)
        p_t.runs[0].font.size = Pt(18)
        p_t.runs[0].font.name = 'Arial'

    if tipo == "examen":
        p_nombre = doc.add_paragraph()
        p_nombre.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p_nombre.add_run("Nombre y apellidos: _________________________________________")
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.name = 'Arial'

    # Pie
    fecha = datetime.now().strftime("%d/%m/%Y")
    pie_izq = autor if autor else f"Generado con IA · {fecha}"
    pie_der = fecha if autor else ""
    for sec in doc.sections:
        añadir_pie(sec, pie_izq, pie_der)

    doc.add_paragraph()

    # Cuerpo
    en_clave = False
    prefijos_re = re.compile(
        r'^(RESUMEN|ESQUEMA|EXAMEN|CUESTIONARIO|PUNTOS CLAVE):\s+', re.IGNORECASE)
    # Para no repetir el título en el cuerpo, registrar qué línea se usó
    lineas_usadas_titulo = set()
    for i, linea in enumerate(lineas):
        l = linea.strip()
        if l and not es_separador(l) and not l.startswith(("✔", "✘")):
            lineas_usadas_titulo.add(i)
            break

    for idx_linea, linea in enumerate(lineas):
        l = linea.strip()
        if es_separador(l): continue
        if l.startswith(("✔ Guardado:", "✘ ERROR")): continue
        if prefijos_re.match(l): continue
        if idx_linea in lineas_usadas_titulo: continue  # saltar línea del título
        if not l:
            doc.add_paragraph(); continue

        if es_clave_respuestas(l):
            en_clave = True
            doc.add_paragraph()
            p = doc.add_heading(l, level=1)
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
                r.font.name = 'Arial'
            continue

        if en_clave:
            p = doc.add_paragraph()
            r = p.add_run(l)
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
            r.font.bold = l.endswith(":") or l.upper() == l
            if r.font.bold:
                r.font.color.rgb = RGBColor(0x5b, 0x4f, 0xcf)
            continue

        if es_cabecera_seccion(l):
            p = doc.add_heading(l, level=1)
            for r in p.runs:
                r.font.color.rgb = RGBColor(0x5b, 0x4f, 0xcf)
                r.font.size = Pt(12)
                r.font.name = 'Arial'
            continue

        if es_pregunta_numerada(l):
            p = doc.add_paragraph(style="List Number")
            m = re.match(r'^\s*\d+[).]\s+(.*)', l)
            r = p.add_run(m.group(1) if m else l)
            r.font.bold = True
            r.font.size = Pt(10.5)
            r.font.name = 'Arial'
            continue

        if es_opcion_multiple(l):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(1)
            r = p.add_run(l)
            r.font.size = Pt(10)
            r.font.name = 'Arial'
            continue

        if l.startswith("► "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(4)
            r = p.add_run("► " + l[2:].strip())
            r.font.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x5b, 0x4f, 0xcf)
            r.font.name = 'Arial'
            continue

        if l.startswith(("•", "- ", "· ")):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(1)
            r = p.add_run(l.lstrip("•-· ").strip())
            r.font.size = Pt(10)
            r.font.name = 'Arial'
            continue

        if l.startswith(("  -", "    -")):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(2)
            r = p.add_run(l.strip().lstrip("- ").strip())
            r.font.size = Pt(9.5)
            r.font.name = 'Arial'
            continue

        p = doc.add_paragraph()
        r = p.add_run(l)
        r.font.size = Pt(10.5)
        r.font.name = 'Arial'

    doc.save(str(ruta_docx))


# ── Conversión CON plantilla ──────────────────────────────────────────────────

def convertir_con_plantilla(ruta_txt: Path, ruta_docx: Path,
                             ruta_plantilla: Path, modulo: str, espad: str):
    """
    Copia la plantilla, actualiza los campos en rojo con módulo y nivel ESPAD,
    y sustituye los marcadores 'Pregunta X.' con el contenido del TXT.
    """
    import shutil
    shutil.copy(str(ruta_plantilla), str(ruta_docx))
    doc = Document(str(ruta_docx))

    # 1. Actualizar campos rojos en la tabla de cabecera
    if doc.tables:
        for row in doc.tables[0].rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        try:
                            if (run.font.color.type and
                                    str(run.font.color.rgb).upper() == "EE0000"):
                                texto_celda = cell.text.upper()
                                if "MÓDULO" in texto_celda or "MODULO" in texto_celda:
                                    run.text = modulo
                                elif "ESPAD" in texto_celda:
                                    run.text = espad
                        except Exception:
                            pass

    # 2. Limpiar contenido del TXT
    lineas = ruta_txt.read_text(encoding='utf-8').splitlines()
    prefijos_re = re.compile(
        r'^(RESUMEN|ESQUEMA|EXAMEN|CUESTIONARIO|PUNTOS CLAVE):\s+', re.IGNORECASE)
    contenido_limpio = [l for l in lineas
                        if not es_separador(l.strip())
                        and not l.strip().startswith(("\u2714", "\u2718"))
                        and not prefijos_re.match(l.strip())]

    # 3. Encontrar inicio (Pregunta 1.) y fin (Calificación total)
    parrafos = doc.paragraphs
    inicio_idx = next((i for i, p in enumerate(parrafos)
                       if re.match(r'^Pregunta\s+1\.', p.text.strip())), None)
    parrafo_cal = next((p for p in parrafos if "Calificación total" in p.text), None)

    if inicio_idx is None:
        inicio_idx = 3  # fallback: tras instrucciones
    if parrafo_cal is None:
        parrafo_cal = parrafos[-1]

    # 4. Eliminar párrafos placeholder
    fin_idx = next((i for i, p in enumerate(parrafos) if "Calificación total" in p.text),
                   len(parrafos))
    for p in list(parrafos[inicio_idx:fin_idx]):
        p._element.getparent().remove(p._element)

    # 5. Insertar contenido nuevo antes de parrafo_cal
    body = doc.element.body
    en_clave = False

    for linea in contenido_limpio:
        l = linea.strip()
        p_new = doc.add_paragraph()
        body.remove(p_new._element)
        parrafo_cal._element.addprevious(p_new._element)

        if not l:
            continue

        if es_clave_respuestas(l):
            en_clave = True
            r = p_new.add_run(l)
            r.font.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
            r.font.name = 'Arial'
            continue

        r = p_new.add_run(l)
        r.font.size = Pt(10.5)
        r.font.name = 'Arial'

        if en_clave:
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
            r.font.bold = l.endswith(":") or (l.upper() == l and len(l) > 3)
        elif es_cabecera_seccion(l):
            r.font.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x5b, 0x4f, 0xcf)
        elif es_pregunta_numerada(l):
            m = re.match(r'^\s*\d+[).]\s+(.*)', l)
            r.text = m.group(1) if m else l
            r.font.bold = True
        elif es_opcion_multiple(l):
            p_new.paragraph_format.left_indent = Cm(1)
        elif l.startswith("\u25ba "):
            r.font.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x5b, 0x4f, 0xcf)
            p_new.paragraph_format.left_indent = Cm(0.5)
        elif l.startswith(("\u2022", "- ", "\u00b7 ")):
            r.text = l.lstrip("\u2022-\u00b7 ").strip()
            p_new.paragraph_format.left_indent = Cm(1)

    doc.save(str(ruta_docx))

def convertir_sin_plantilla(ruta_txt: Path, ruta_docx: Path, autor: str = ""):
    doc = Document()

    # Fuente Arial global
    for style in doc.styles:
        try: style.font.name = 'Arial'
        except Exception: pass
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10.5)

    # Márgenes
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3)
        sec.right_margin  = Cm(3)

    lineas = ruta_txt.read_text(encoding='utf-8').splitlines()
    tipo = detectar_tipo(lineas)

    # Título
    titulo_doc = ruta_txt.stem.replace("_", " ").title()
    for linea in lineas:
        l = linea.strip()
        if l and not es_separador(l) and not l.startswith(("✔", "✘")):
            for pref in ["RESUMEN:", "ESQUEMA:", "EXAMEN:", "CUESTIONARIO:", "PUNTOS CLAVE:"]:
                if l.upper().startswith(pref):
                    l = l[len(pref):].strip()
            if l:
                titulo_doc = l
                break

    if tipo == "examen":
        p_t = doc.add_heading("Examen de Comunicación", level=0)
    else:
        p_t = doc.add_heading(titulo_doc, level=0)
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_t.runs:
        p_t.runs[0].font.color.rgb = RGBColor(0x5b, 0x4f, 0xcf)
        p_t.runs[0].font.size = Pt(18)
        p_t.runs[0].font.name = 'Arial'

    if tipo == "examen":
        p_nombre = doc.add_paragraph()
        p_nombre.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p_nombre.add_run("Nombre y apellidos: _________________________________________")
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.name = 'Arial'

    # Pie
    fecha = datetime.now().strftime("%d/%m/%Y")
    pie_izq = autor if autor else f"Generado con IA · {fecha}"
    pie_der = fecha if autor else ""
    for sec in doc.sections:
        añadir_pie(sec, pie_izq, pie_der)

    doc.add_paragraph()

    # Cuerpo
    en_clave = False
    prefijos_re = re.compile(
        r'^(RESUMEN|ESQUEMA|EXAMEN|CUESTIONARIO|PUNTOS CLAVE):\s+', re.IGNORECASE)
    # Para no repetir el título en el cuerpo, registrar qué línea se usó
    lineas_usadas_titulo = set()
    for i, linea in enumerate(lineas):
        l = linea.strip()
        if l and not es_separador(l) and not l.startswith(("✔", "✘")):
            lineas_usadas_titulo.add(i)
            break

    for idx_linea, linea in enumerate(lineas):
        l = linea.strip()
        if es_separador(l): continue
        if l.startswith(("✔ Guardado:", "✘ ERROR")): continue
        if prefijos_re.match(l): continue
        if idx_linea in lineas_usadas_titulo: continue  # saltar línea del título
        if not l:
            doc.add_paragraph(); continue

        if es_clave_respuestas(l):
            en_clave = True
            doc.add_paragraph()
            p = doc.add_heading(l, level=1)
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
                r.font.name = 'Arial'
            continue

        if en_clave:
            p = doc.add_paragraph()
            r = p.add_run(l)
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
            r.font.bold = l.endswith(":") or l.upper() == l
            if r.font.bold:
                r.font.color.rgb = RGBColor(0x5b, 0x4f, 0xcf)
            continue

        if es_cabecera_seccion(l):
            p = doc.add_heading(l, level=1)
            for r in p.runs:
                r.font.color.rgb = RGBColor(0x5b, 0x4f, 0xcf)
                r.font.size = Pt(12)
                r.font.name = 'Arial'
            continue

        if es_pregunta_numerada(l):
            p = doc.add_paragraph(style="List Number")
            m = re.match(r'^\s*\d+[).]\s+(.*)', l)
            r = p.add_run(m.group(1) if m else l)
            r.font.bold = True
            r.font.size = Pt(10.5)
            r.font.name = 'Arial'
            continue

        if es_opcion_multiple(l):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(1)
            r = p.add_run(l)
            r.font.size = Pt(10)
            r.font.name = 'Arial'
            continue

        if l.startswith("► "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(4)
            r = p.add_run("► " + l[2:].strip())
            r.font.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x5b, 0x4f, 0xcf)
            r.font.name = 'Arial'
            continue

        if l.startswith(("•", "- ", "· ")):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(1)
            r = p.add_run(l.lstrip("•-· ").strip())
            r.font.size = Pt(10)
            r.font.name = 'Arial'
            continue

        if l.startswith(("  -", "    -")):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(2)
            r = p.add_run(l.strip().lstrip("- ").strip())
            r.font.size = Pt(9.5)
            r.font.name = 'Arial'
            continue

        p = doc.add_paragraph()
        r = p.add_run(l)
        r.font.size = Pt(10.5)
        r.font.name = 'Arial'

    doc.save(str(ruta_docx))


# ── Conversión CON plantilla ──────────────────────────────────────────────────

def convertir_con_plantilla(ruta_txt: Path, ruta_docx: Path,
                             ruta_plantilla: Path, modulo: str, espad: str):
    """
    Copia la plantilla, actualiza los campos en rojo con módulo y nivel ESPAD,
    y sustituye los marcadores 'Pregunta X.' con el contenido del TXT.
    """
    import shutil
    shutil.copy(str(ruta_plantilla), str(ruta_docx))
    doc = Document(str(ruta_docx))

    # 1. Actualizar campos rojos en la tabla de cabecera
    if doc.tables:
        tabla = doc.tables[0]
        for row in tabla.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        # Módulo (rojo en col 1)
                        if (run.font.color.type and
                                str(run.font.color.rgb).upper() == "EE0000"):
                            texto_celda = cell.text.upper()
                            if "MÓDULO" in texto_celda or "MODULO" in texto_celda:
                                run.text = modulo
                            elif "ESPAD" in texto_celda:
                                run.text = espad

    # 2. Extraer preguntas del TXT
    lineas = ruta_txt.read_text(encoding='utf-8').splitlines()
    prefijos_re = re.compile(
        r'^(RESUMEN|ESQUEMA|EXAMEN|CUESTIONARIO|PUNTOS CLAVE):\s+', re.IGNORECASE)

    contenido_limpio = []
    for l in lineas:
        ls = l.strip()
        if es_separador(ls): continue
        if ls.startswith(("✔ Guardado:", "✘ ERROR")): continue
        if prefijos_re.match(ls): continue
        contenido_limpio.append(l)

    contenido_txt = "\n".join(contenido_limpio).strip()

    # 3. Localizar y reemplazar marcadores "Pregunta X." en el documento
    # Busca el primer párrafo "Pregunta 1." y reemplaza desde ahí
    inicio_idx = None
    for i, p in enumerate(doc.paragraphs):
        if re.match(r'^Pregunta\s+1\.', p.text.strip()):
            inicio_idx = i
            break

    if inicio_idx is not None:
        # Eliminar todos los párrafos desde Pregunta 1 hasta el pie (Calificación)
        fin_idx = len(doc.paragraphs)
        for i, p in enumerate(doc.paragraphs):
            if "Calificación total" in p.text:
                fin_idx = i
                break

        # Eliminar párrafos del cuerpo del examen (de fin a inicio, para no romper índices)
        parrafos_a_eliminar = doc.paragraphs[inicio_idx:fin_idx]
        for p in parrafos_a_eliminar:
            p._element.getparent().remove(p._element)

        # Insertar contenido del TXT antes del párrafo de Calificación
        # Usamos el elemento padre del documento
        body = doc.element.body

        # Punto de inserción: antes del párrafo de Calificación
        parrafo_cal = None
        for p in doc.paragraphs:
            if "Calificación total" in p.text:
                parrafo_cal = p
                break

        def nuevo_parrafo_antes(doc_obj, ref_p, texto, bold=False, size=10.5,
                                color=None, indent_cm=0, estilo='Normal'):
            from docx.oxml.ns import qn
            p_new = doc_obj.add_paragraph(style=estilo)
            r = p_new.add_run(texto)
            r.font.size = Pt(size)
            r.font.name = 'Arial'
            r.font.bold = bold
            if color:
                r.font.color.rgb = color
            if indent_cm:
                p_new.paragraph_format.left_indent = Cm(indent_cm)
            if ref_p is not None:
                ref_p._element.addprevious(p_new._element)
                # Remove from end (add_paragraph adds to end)
                body.remove(p_new._element)
                ref_p._element.addprevious(p_new._element)
            return p_new

        en_clave = False
        for linea in contenido_limpio:
            l = linea.strip()
            if not l:
                nuevo_parrafo_antes(doc, parrafo_cal, "")
                continue

            if es_clave_respuestas(l):
                en_clave = True
                nuevo_parrafo_antes(doc, parrafo_cal, "")
                nuevo_parrafo_antes(doc, parrafo_cal, l, bold=True, size=12,
                                    color=RGBColor(0xc0, 0x39, 0x2b))
                continue

            if en_clave:
                nuevo_parrafo_antes(doc, parrafo_cal, l, size=10,
                                    bold=(l.endswith(":") or l.upper() == l))
                continue

            if es_cabecera_seccion(l):
                nuevo_parrafo_antes(doc, parrafo_cal, l, bold=True, size=12,
                                    color=RGBColor(0x5b, 0x4f, 0xcf))
                continue

            if es_pregunta_numerada(l):
                m = re.match(r'^\s*\d+[).]\s+(.*)', l)
                nuevo_parrafo_antes(doc, parrafo_cal,
                                    m.group(1) if m else l,
                                    bold=True, size=10.5, estilo='Body Text')
                continue

            if es_opcion_multiple(l):
                nuevo_parrafo_antes(doc, parrafo_cal, l, size=10, indent_cm=1)
                continue

            if l.startswith("► "):
                nuevo_parrafo_antes(doc, parrafo_cal, l, bold=True, size=11,
                                    color=RGBColor(0x5b, 0x4f, 0xcf), indent_cm=0.5)
                continue

            if l.startswith(("•", "- ", "· ")):
                nuevo_parrafo_antes(doc, parrafo_cal,
                                    l.lstrip("•-· ").strip(), size=10, indent_cm=1)
                continue

            nuevo_parrafo_antes(doc, parrafo_cal, l, size=10.5)

    doc.save(str(ruta_docx))


# ── GUI ───────────────────────────────────────────────────────────────────────

class ConvertidorApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Convertidor TXT → DOCX")
        self.configure(bg=BG)
        self.resizable(True, True)
        self._construir_ui()
        self._centrar_ventana(660, 560)

    def _centrar_ventana(self, w, h):
        self.update_idletasks()
        x = (self.winfo_screenwidth() - w) // 2
        self.geometry(f"{w}x{h}+{x}+5")
        self.minsize(520, 460)

    def _construir_ui(self):
        ttk.Style().theme_use("clam")

        # Cabecera
        cab = tk.Frame(self, bg=ACENTO, pady=10)
        cab.pack(fill="x")
        tk.Label(cab, text="📄 Convertidor TXT → DOCX",
                 font=F_TITULO, bg=ACENTO, fg="white").pack(side="left", padx=18)
        tk.Label(cab, text="Complemento a resúmenes y exámenes generados con Groq",
                 font=F_LABEL, bg=ACENTO, fg="#ccc8ff").pack(side="right", padx=18)

        cuerpo = tk.Frame(self, bg=BG, padx=20, pady=12)
        cuerpo.pack(fill="both", expand=True)

        # Lista de archivos TXT
        self._lbl(cuerpo, "📋  Archivos TXT a convertir")
        fr_txt = tk.Frame(cuerpo, bg=BORDE, bd=1, relief="solid")
        fr_txt.pack(fill="both", expand=True, pady=(0, 6))
        self.lst_archivos = tk.Listbox(fr_txt, font=F_MONO, bg=ENTRADA, fg=TEXTO,
                                       selectbackground=ACENTO, selectforeground="white",
                                       relief="flat", bd=6, height=5)
        sb = ttk.Scrollbar(fr_txt, orient="vertical", command=self.lst_archivos.yview)
        self.lst_archivos.configure(yscrollcommand=sb.set)
        self.lst_archivos.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        self.lst_archivos.bind("<<ListboxSelect>>", self._al_seleccionar)

        # Botones lista
        fr_bts = tk.Frame(cuerpo, bg=BG)
        fr_bts.pack(fill="x", pady=(0, 8))
        tk.Button(fr_bts, text="＋ Agregar TXT", font=F_NORM,
                  bg=PANEL, fg=TEXTO, relief="flat", cursor="hand2",
                  padx=10, pady=4, command=self._agregar).pack(side="left")
        tk.Button(fr_bts, text="✖ Quitar", font=F_NORM,
                  bg=PANEL, fg=SUBTEXTO, relief="flat", cursor="hand2",
                  padx=10, pady=4, command=self._quitar).pack(side="left", padx=6)

        # Plantilla (opcional)
        self._lbl(cuerpo, "📑  Plantilla .docx (opcional — solo para exámenes)")
        fr_pl = tk.Frame(cuerpo, bg=BG)
        fr_pl.pack(fill="x", pady=(0, 6))
        self.var_plantilla = tk.StringVar()
        tk.Entry(fr_pl, textvariable=self.var_plantilla, font=F_NORM,
                 bg=ENTRADA, fg=TEXTO, relief="solid", bd=1).pack(
                     side="left", fill="x", expand=True)
        tk.Button(fr_pl, text="…", font=F_NORM, bg=ACENTO, fg="white",
                  relief="flat", cursor="hand2",
                  command=self._elegir_plantilla).pack(side="left", padx=(3, 0))
        tk.Button(fr_pl, text="✖", font=F_NORM, bg=PANEL, fg=SUBTEXTO,
                  relief="flat", cursor="hand2",
                  command=lambda: self.var_plantilla.set("")).pack(side="left", padx=(2, 0))

        # Campos en rojo de la plantilla (Módulo y ESPAD)
        self.frame_meta = tk.Frame(cuerpo, bg=PANEL, padx=8, pady=6,
                                   relief="solid", bd=1)
        tk.Label(self.frame_meta,
                 text="Campos de la plantilla (extraídos del nombre del TXT — editables):",
                 font=F_BOLD, bg=PANEL, fg=SUBTEXTO).pack(anchor="w", pady=(0, 4))
        fr_m = tk.Frame(self.frame_meta, bg=PANEL)
        fr_m.pack(fill="x", pady=2)
        tk.Label(fr_m, text="Módulo:", font=F_NORM, bg=PANEL,
                 fg=TEXTO, width=10, anchor="w").pack(side="left")
        self.var_modulo = tk.StringVar()
        tk.Entry(fr_m, textvariable=self.var_modulo, font=F_NORM,
                 bg=ENTRADA, fg=TEXTO, relief="solid", bd=1).pack(
                     side="left", fill="x", expand=True)
        fr_e = tk.Frame(self.frame_meta, bg=PANEL)
        fr_e.pack(fill="x", pady=2)
        tk.Label(fr_e, text="ESPAD:", font=F_NORM, bg=PANEL,
                 fg=TEXTO, width=10, anchor="w").pack(side="left")
        self.var_espad = tk.StringVar()
        tk.Entry(fr_e, textvariable=self.var_espad, font=F_NORM,
                 bg=ENTRADA, fg=TEXTO, relief="solid", bd=1,
                 width=10).pack(side="left")

        # Carpeta destino
        self._lbl(cuerpo, "📁  Carpeta de destino")
        fr_carp = tk.Frame(cuerpo, bg=BG)
        fr_carp.pack(fill="x", pady=(0, 4))
        self.var_carpeta = tk.StringVar(value=str(Path.home() / "Resumenes"))
        tk.Entry(fr_carp, textvariable=self.var_carpeta, font=F_NORM,
                 bg=ENTRADA, fg=TEXTO, relief="solid", bd=1).pack(
                     side="left", fill="x", expand=True)
        tk.Button(fr_carp, text="…", font=F_NORM, bg=ACENTO, fg="white",
                  relief="flat", cursor="hand2",
                  command=self._elegir_carpeta).pack(side="left", padx=(3, 0))

        # Autor (sin plantilla)
        self._lbl(cuerpo, "✏️  Autor / Centro (pie de página, sin plantilla)")
        self.var_autor = tk.StringVar()
        tk.Entry(cuerpo, textvariable=self.var_autor, font=F_NORM,
                 bg=ENTRADA, fg=TEXTO, relief="solid", bd=1).pack(
                     fill="x", pady=(0, 10))

        # Botones acción
        fr_acc = tk.Frame(cuerpo, bg=BG)
        fr_acc.pack(fill="x")
        tk.Button(fr_acc, text="📂  Abrir carpeta", font=F_LABEL,
                  bg=PANEL, fg=ACENTO2, relief="flat", cursor="hand2",
                  pady=5, command=self._abrir_carpeta).pack(side="left")
        self.btn_conv = tk.Button(
            fr_acc, text="▶  CONVERTIR",
            font=("Segoe UI", 10, "bold"), bg=ACENTO, fg="white",
            relief="flat", cursor="hand2", padx=16, pady=5,
            command=self._convertir)
        self.btn_conv.pack(side="right")

        # Barra estado
        barra = tk.Frame(self, bg=PANEL)
        barra.pack(fill="x", side="bottom")
        self.var_estado = tk.StringVar(value="Listo.")
        tk.Label(barra, textvariable=self.var_estado,
                 font=F_LABEL, bg=PANEL, fg=SUBTEXTO,
                 anchor="w").pack(side="left", padx=12, pady=5)

        # Mostrar/ocultar frame_meta según plantilla
        self.var_plantilla.trace_add("write", self._toggle_meta)

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _lbl(self, padre, texto):
        tk.Label(padre, text=texto, font=F_BOLD,
                 bg=BG, fg=SUBTEXTO).pack(anchor="w", pady=(6, 2))

    def _toggle_meta(self, *_):
        if self.var_plantilla.get().strip():
            self.frame_meta.pack(fill="x", pady=(0, 8),
                                 before=self._widget_after_plantilla())
        else:
            self.frame_meta.pack_forget()

    def _widget_after_plantilla(self):
        # Devuelve el widget "Carpeta de destino" label para insertar antes
        for w in self.frame_meta.master.pack_slaves():
            if isinstance(w, tk.Label) and "Carpeta" in str(w.cget("text")):
                return w
        return None

    def _al_seleccionar(self, _event=None):
        """Auto-rellena Módulo y ESPAD al seleccionar un archivo."""
        sel = self.lst_archivos.curselection()
        if not sel: return
        nombre = Path(self.lst_archivos.get(sel[0])).stem
        meta = extraer_meta_de_nombre(nombre)
        if meta['modulo']: self.var_modulo.set(meta['modulo'])
        if meta['espad']:  self.var_espad.set(meta['espad'])

    def _agregar(self):
        archivos = filedialog.askopenfilenames(
            title="Selecciona archivos TXT",
            filetypes=[("Texto", "*.txt"), ("Todos", "*.*")])
        for r in archivos:
            if r not in self.lst_archivos.get(0, "end"):
                self.lst_archivos.insert("end", r)
        # Auto-seleccionar primero
        if self.lst_archivos.size() > 0:
            self.lst_archivos.selection_set(0)
            self._al_seleccionar()

    def _quitar(self):
        for i in reversed(self.lst_archivos.curselection()):
            self.lst_archivos.delete(i)

    def _elegir_plantilla(self):
        f = filedialog.askopenfilename(
            title="Selecciona plantilla Word",
            filetypes=[("Word", "*.docx"), ("Todos", "*.*")])
        if f: self.var_plantilla.set(f)

    def _elegir_carpeta(self):
        d = filedialog.askdirectory()
        if d: self.var_carpeta.set(d)

    def _abrir_carpeta(self):
        carpeta = self.var_carpeta.get().strip()
        if not carpeta: return
        Path(carpeta).mkdir(parents=True, exist_ok=True)
        if sys.platform   == "win32":  os.startfile(carpeta)
        elif sys.platform == "darwin": subprocess.Popen(["open", carpeta])
        else:                          subprocess.Popen(["xdg-open", carpeta])

    def _convertir(self):
        archivos = list(self.lst_archivos.get(0, "end"))
        if not archivos:
            self.var_estado.set("⚠ Agrega al menos un archivo TXT.")
            return
        carpeta     = Path(self.var_carpeta.get().strip())
        plantilla   = self.var_plantilla.get().strip()
        autor       = self.var_autor.get().strip()
        modulo      = self.var_modulo.get().strip()
        espad       = self.var_espad.get().strip()
        carpeta.mkdir(parents=True, exist_ok=True)

        ok = err = 0
        for ruta_str in archivos:
            ruta = Path(ruta_str)
            ruta_docx = carpeta / (ruta.stem + ".docx")
            # Si hay varios archivos, auto-extraer meta para cada uno
            meta = extraer_meta_de_nombre(ruta.stem)
            mod_usar  = modulo or meta['modulo']
            esp_usar  = espad  or meta['espad']
            try:
                self.var_estado.set(f"Convirtiendo: {ruta.name}…")
                self.update()
                if plantilla and Path(plantilla).exists():
                    self.var_estado.set(f"Usando plantilla para: {ruta.name}…")
                    self.update()
                    convertir_con_plantilla(ruta, ruta_docx,
                                            Path(plantilla), mod_usar, esp_usar)
                else:
                    convertir_sin_plantilla(ruta, ruta_docx, autor)
                if ruta_docx.exists():
                    ok += 1
                else:
                    raise Exception(f"El archivo no se creó: {ruta_docx}")
            except Exception as e:
                err += 1
                import traceback
                msg = traceback.format_exc()
                self.var_estado.set(f"✘ Error: {e}")
                # Mostrar error completo en ventana emergente
                import tkinter.messagebox as mb
                mb.showerror("Error en conversión", f"{ruta.name}:\n\n{msg}")

        import tkinter.messagebox as mb
        if err == 0:
            msg = str(ok) + " archivo(s) convertido(s) correctamente.\n\nCarpeta: " + str(carpeta)
            self.var_estado.set("Listo. " + str(ok) + " archivo(s) convertidos.")
            mb.showinfo("Conversion completada", msg)
        else:
            msg = str(ok) + " convertido(s). " + str(err) + " con error.\n\nCarpeta: " + str(carpeta)
            self.var_estado.set("Listo con errores.")
            mb.showwarning("Conversion con errores", msg)


if __name__ == "__main__":
    app = ConvertidorApp()
    app.mainloop()
